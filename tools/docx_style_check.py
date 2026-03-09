from __future__ import annotations

import argparse
import posixpath
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path, PurePosixPath
from typing import Iterable
from xml.etree import ElementTree as ET

from docx import Document

TARGET_PAGE_W_MM = 210.0
TARGET_PAGE_H_MM = 297.0
TARGET_MARGIN_LEFT_MM = 30.0
TARGET_MARGIN_RIGHT_MM = 10.0
TARGET_MARGIN_TOP_MM = 20.0
TARGET_MARGIN_BOTTOM_MM = 20.0

TARGET_NORMAL_FONT = "Times New Roman"
TARGET_NORMAL_SIZE_PT = 14.0
TARGET_LINE_SPACING = 1.5
TARGET_FIRST_LINE_INDENT_CM = 1.25

TOL_MM = 0.6
TOL_PT = 0.6
TOL_CM = 0.1


@dataclass(frozen=True)
class DocxIssue:
    path: Path
    message: str


def _approx(a: float, b: float, tol: float) -> bool:
    return abs(a - b) <= tol


def _iter_docx_files(root: Path) -> Iterable[Path]:
    if root.is_file() and root.suffix.lower() == ".docx":
        yield root
        return
    if not root.exists():
        return
    for p in sorted(root.rglob("*.docx")):
        if p.is_file():
            yield p


def _zip_read(z: zipfile.ZipFile, name: str) -> bytes:
    return z.read(name)


def _norm_word_target(target: str) -> str:
    # Targets in *_rels are relative to the parent part (word/document.xml).
    # Normalize ".." segments to keep lookups stable.
    base = PurePosixPath("word")
    return posixpath.normpath(str(base / PurePosixPath(target)))


def _docx_rels_map(z: zipfile.ZipFile) -> dict[str, str]:
    """
    Map rId -> target part path, e.g. "rId10" -> "word/footer2.xml".
    """
    rels_name = "word/_rels/document.xml.rels"
    if rels_name not in z.namelist():
        return {}
    data = _zip_read(z, rels_name)
    try:
        root = ET.fromstring(data)
    except Exception:
        return {}

    rel_ns = "{http://schemas.openxmlformats.org/package/2006/relationships}"
    out: dict[str, str] = {}
    for rel in root.findall(f"{rel_ns}Relationship"):
        rid = rel.attrib.get("Id") or ""
        target = rel.attrib.get("Target") or ""
        if rid and target:
            out[rid] = _norm_word_target(target)
    return out


def _docx_part_has_field(part_xml: bytes, field: str) -> bool:
    # Field instructions live in w:instrText.
    low = part_xml.lower()
    return b"instrtext" in low and field.lower().encode("utf-8") in low


def _docx_has_wp_anchor(z: zipfile.ZipFile) -> bool:
    for name in z.namelist():
        if not (name.startswith("word/") and name.endswith(".xml")):
            continue
        data = _zip_read(z, name)
        if b"wp:anchor" in data:
            return True
    return False


def _check_footer_page_number(docx_path: Path) -> list[str]:
    """
    Enforce:
    - section has different first page footer enabled
    - default footer has PAGE field
    - first-page footer does NOT have PAGE field
    """
    errors: list[str] = []
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            rels = _docx_rels_map(z)
            if "word/document.xml" not in z.namelist():
                return ["Missing word/document.xml"]
            doc_xml = _zip_read(z, "word/document.xml")
            try:
                root = ET.fromstring(doc_xml)
            except Exception as exc:
                return [f"Failed to parse word/document.xml: {exc}"]

            w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
            r_ns = (
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
            )

            # Usually one section; check the first encountered sectPr.
            sect = root.find(f".//{w_ns}sectPr")
            if sect is None:
                return ["Missing w:sectPr in document.xml"]

            refs: dict[str, str] = {}
            for fr in sect.findall(f"{w_ns}footerReference"):
                typ = fr.attrib.get(f"{w_ns}type") or ""
                rid = fr.attrib.get(f"{r_ns}id") or ""
                if typ and rid:
                    refs[typ] = rid

            if "default" not in refs:
                errors.append('Missing default footer reference (w:type="default")')
                return errors
            if "first" not in refs:
                errors.append('Missing first-page footer reference (w:type="first")')
                return errors

            def_part = rels.get(refs["default"], "")
            first_part = rels.get(refs["first"], "")
            if not def_part or def_part not in z.namelist():
                errors.append(
                    f"Default footer part not found for {refs['default']}: {def_part}"
                )
                return errors
            if not first_part or first_part not in z.namelist():
                errors.append(
                    f"First-page footer part not found for {refs['first']}: {first_part}"
                )
                return errors

            def_xml = _zip_read(z, def_part)
            first_xml = _zip_read(z, first_part)
            if not _docx_part_has_field(def_xml, "PAGE"):
                errors.append("Default footer does not contain PAGE field")
            if _docx_part_has_field(first_xml, "PAGE"):
                errors.append(
                    "First-page footer contains PAGE field (should be empty on title page)"
                )
    except zipfile.BadZipFile:
        errors.append("Not a valid DOCX/ZIP")
    except Exception as exc:  # noqa: BLE001
        errors.append(f"Footer check failed: {exc}")
    return errors


def _check_docx(path: Path) -> list[str]:
    errors: list[str] = []
    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        return [f"Failed to open: {exc}"]

    if not doc.sections:
        return ["No sections found"]
    sec = doc.sections[0]

    # Page size
    if not _approx(sec.page_width.mm, TARGET_PAGE_W_MM, TOL_MM) or not _approx(
        sec.page_height.mm, TARGET_PAGE_H_MM, TOL_MM
    ):
        errors.append(
            f"Page size not A4: {sec.page_width.mm:.2f}x{sec.page_height.mm:.2f} mm"
        )

    # Margins
    if not _approx(sec.left_margin.mm, TARGET_MARGIN_LEFT_MM, TOL_MM):
        errors.append(
            f"Left margin: {sec.left_margin.mm:.2f} mm (expected {TARGET_MARGIN_LEFT_MM} mm)"
        )
    if not _approx(sec.right_margin.mm, TARGET_MARGIN_RIGHT_MM, TOL_MM):
        errors.append(
            f"Right margin: {sec.right_margin.mm:.2f} mm (expected {TARGET_MARGIN_RIGHT_MM} mm)"
        )
    if not _approx(sec.top_margin.mm, TARGET_MARGIN_TOP_MM, TOL_MM):
        errors.append(
            f"Top margin: {sec.top_margin.mm:.2f} mm (expected {TARGET_MARGIN_TOP_MM} mm)"
        )
    if not _approx(sec.bottom_margin.mm, TARGET_MARGIN_BOTTOM_MM, TOL_MM):
        errors.append(
            f"Bottom margin: {sec.bottom_margin.mm:.2f} mm (expected {TARGET_MARGIN_BOTTOM_MM} mm)"
        )

    # Normal style
    try:
        normal = doc.styles["Normal"]
        if (normal.font.name or "").strip() != TARGET_NORMAL_FONT:
            errors.append(
                f"Normal font: '{normal.font.name}' (expected '{TARGET_NORMAL_FONT}')"
            )
        if normal.font.size is None or not _approx(
            float(normal.font.size.pt), TARGET_NORMAL_SIZE_PT, TOL_PT
        ):
            got = float(normal.font.size.pt) if normal.font.size else None
            errors.append(
                f"Normal font size: {got} pt (expected {TARGET_NORMAL_SIZE_PT} pt)"
            )
        pf = normal.paragraph_format
        if pf.line_spacing is None or not _approx(
            float(pf.line_spacing), TARGET_LINE_SPACING, 0.05
        ):
            errors.append(
                f"Normal line spacing: {pf.line_spacing} (expected {TARGET_LINE_SPACING})"
            )
        if pf.first_line_indent is None or not _approx(
            float(pf.first_line_indent.cm), TARGET_FIRST_LINE_INDENT_CM, TOL_CM
        ):
            got = float(pf.first_line_indent.cm) if pf.first_line_indent else None
            errors.append(
                f"Normal first-line indent: {got} cm (expected {TARGET_FIRST_LINE_INDENT_CM} cm)"
            )
    except Exception as exc:  # noqa: BLE001
        errors.append(f"Normal style check failed: {exc}")

    # First page numbering rule
    if not bool(sec.different_first_page_header_footer):
        errors.append(
            "different_first_page_header_footer is False (title page should have empty footer)"
        )

    # Header text must be empty (no visible text)
    header_txt = "\n".join((p.text or "") for p in sec.header.paragraphs).strip()
    first_header_txt = "\n".join(
        (p.text or "") for p in sec.first_page_header.paragraphs
    ).strip()
    if header_txt:
        errors.append(f"Non-empty header text: '{header_txt[:120]}'")
    if first_header_txt:
        errors.append(f"Non-empty first-page header text: '{first_header_txt[:120]}'")

    # Pictures must fit within text width
    try:
        avail_emu = (
            int(sec.page_width.emu)
            - int(sec.left_margin.emu)
            - int(sec.right_margin.emu)
        )
        for i, shp in enumerate(doc.inline_shapes, start=1):
            if int(shp.width) > avail_emu:
                errors.append(
                    f"InlineShape #{i} width {int(shp.width)} EMU exceeds text width {avail_emu} EMU"
                )
    except Exception:
        pass

    # Tables should not be autofit
    for i, t in enumerate(doc.tables, start=1):
        try:
            if bool(t.autofit):
                errors.append(f"Table #{i}: autofit is ON (should be OFF)")
        except Exception:
            continue

    # Low-level XML checks (floating objects, footer PAGE)
    try:
        with zipfile.ZipFile(path, "r") as z:
            if _docx_has_wp_anchor(z):
                errors.append("Floating drawing found (wp:anchor)")
    except Exception as exc:  # noqa: BLE001
        errors.append(f"ZIP scan failed: {exc}")

    errors.extend(_check_footer_page_number(path))
    return errors


def main(argv: list[str]) -> int:
    ap = argparse.ArgumentParser(
        description="DOCX style checker for KP_2025 norm-control constraints."
    )
    ap.add_argument(
        "--root",
        default=None,
        help="Root dir or a .docx file to check (default: repo output folder)",
    )
    ap.add_argument(
        "--out",
        default=None,
        help="Write report to this path (default: style_report.txt under root)",
    )
    args = ap.parse_args(argv)

    repo_root = Path(__file__).resolve()
    for cand in (repo_root.parent, *repo_root.parents):
        if (cand / ".git").exists():
            repo_root = cand
            break
    default_root = repo_root / "КП_2025_EyeGate_Mantrap_Шаповалов_ИУ8-74"
    root = Path(args.root).resolve() if args.root else default_root

    out_path = (
        Path(args.out).resolve()
        if args.out
        else (
            root / "style_report.txt"
            if root.is_dir()
            else root.with_suffix(".style_report.txt")
        )
    )
    out_path.parent.mkdir(parents=True, exist_ok=True)

    docx_files = list(_iter_docx_files(root))
    issues: list[DocxIssue] = []
    for p in docx_files:
        errs = _check_docx(p)
        for e in errs:
            issues.append(DocxIssue(path=p, message=e))

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    lines: list[str] = []
    lines.append(f"Style check report: {now}")
    lines.append(f"Root: {root}")
    lines.append(f"DOCX files: {len(docx_files)}")
    lines.append("")

    if not docx_files:
        lines.append("No .docx files found.")
    else:
        by_file: dict[Path, list[str]] = {}
        for it in issues:
            by_file.setdefault(it.path, []).append(it.message)
        for p in docx_files:
            errs = by_file.get(p, [])
            status = "PASS" if not errs else "FAIL"
            rel = (
                p.relative_to(root)
                if root.is_dir() and p.is_absolute() and str(p).startswith(str(root))
                else p
            )
            lines.append(f"[{status}] {rel}")
            for e in errs:
                lines.append(f" - {e}")
            lines.append("")

    fail_files = {i.path for i in issues}
    lines.append(
        f"Summary: PASS={len(docx_files) - len(fail_files)}, FAIL={len(fail_files)}"
    )
    out_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")

    return 0 if not issues else 2


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
