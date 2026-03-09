from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _strip_inline_md(text: str) -> str:
    # Minimal inline markdown cleanup to avoid raw ** / `...` in DOCX output.
    out = text
    out = re.sub(r"\*\*(.+?)\*\*", r"\1", out)
    out = re.sub(r"__(.+?)__", r"\1", out)
    out = re.sub(r"`([^`]+)`", r"\1", out)
    # Links: [text](url) -> text (url)
    out = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", out)
    return out


def _is_table_line(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    if not _is_table_line(s):
        return False
    inner = s.strip("|").strip()
    # Allow :---: alignment markers
    parts = [p.strip() for p in inner.split("|")]
    if not parts:
        return False
    return all(re.fullmatch(r":?-{3,}:?", p or "") is not None for p in parts)


def _split_table_row(line: str) -> list[str]:
    inner = line.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


def _add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph("")
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.0


@dataclass
class HeadingNumbers:
    h1: int = 0
    h2: int = 0
    h3: int = 0

    def next(self, level: int) -> str:
        if level == 1:
            self.h1 += 1
            self.h2 = 0
            self.h3 = 0
            return f"{self.h1}"
        if level == 2:
            self.h2 += 1
            self.h3 = 0
            return f"{self.h1}.{self.h2}"
        if level == 3:
            self.h3 += 1
            return f"{self.h1}.{self.h2}.{self.h3}"
        raise ValueError("Unsupported heading level")


def add_numbered_heading(
    doc: Document, nums: HeadingNumbers, title: str, level: int
) -> None:
    prefix = nums.next(level)
    style = f"Heading {level}"
    p = doc.add_paragraph(f"{prefix} {title}", style=style)
    p.paragraph_format.first_line_indent = None


def _add_bullet(doc: Document, text: str, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    try:
        doc.styles[style]
        p = doc.add_paragraph(_strip_inline_md(text), style=style)
    except KeyError:
        prefix = "- " if numbered else "• "
        p = doc.add_paragraph(prefix + _strip_inline_md(text), style="Normal")
    p.paragraph_format.first_line_indent = None


def append_markdown(
    doc: Document,
    nums: HeadingNumbers,
    md_text: str,
    base_heading_level: int = 2,
    max_heading_level: int = 3,
    drop_first_heading: bool = False,
    heading_prefix: Optional[str] = None,
) -> None:
    """
    Minimal markdown -> docx importer:
    - headings (#/##/###)
    - bullet/numbered lists
    - fenced code blocks
    - simple tables

    Headings are converted into numbered Heading styles using `nums`.
    """
    lines = md_text.splitlines()
    in_code = False
    code_fence = ""
    pending_table: list[str] = []
    skipped_first_heading = False

    def flush_table() -> None:
        nonlocal pending_table
        if not pending_table:
            return
        # table requires at least header + sep
        if len(pending_table) < 2 or not _is_table_sep(pending_table[1]):
            # fallback: render lines as text
            for ln in pending_table:
                doc.add_paragraph(ln.strip(), style="Normal")
            pending_table = []
            return
        headers = _split_table_row(pending_table[0])
        rows = []
        for ln in pending_table[2:]:
            if not _is_table_line(ln):
                continue
            rows.append(_split_table_row(ln))
        cols = len(headers)
        table = doc.add_table(rows=1, cols=cols)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for r in rows:
            cells = table.add_row().cells
            for i in range(cols):
                cells[i].text = r[i] if i < len(r) else ""

        # Widths and formatting (A4: 210mm - (30mm+10mm) = 170mm = 17.0cm).
        try:
            col_w = Cm(17.0 / max(1, cols))
            for tr in table.rows:
                for tc in tr.cells:
                    tc.width = col_w
                    tc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for p in tc.paragraphs:
                        p.paragraph_format.first_line_indent = None
                        p.paragraph_format.line_spacing = 1.0
                        for run in p.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)
        except Exception:
            pass
        pending_table = []

    for raw in lines:
        line = raw.rstrip("\n")

        # Code fences
        if line.strip().startswith("```"):
            fence = line.strip()
            if not in_code:
                in_code = True
                code_fence = fence
                flush_table()
                continue
            # closing fence
            if in_code:
                in_code = False
                code_fence = ""
                continue

        if in_code:
            _add_code_paragraph(doc, line)
            continue

        # Tables (markdown pipe tables)
        if _is_table_line(line):
            pending_table.append(line)
            continue
        if pending_table:
            # end of table
            flush_table()

        s = line.strip()
        if not s:
            doc.add_paragraph("", style="Normal")
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            hashes = len(m.group(1))
            title = m.group(2).strip()
            if heading_prefix:
                title = f"{heading_prefix}: {title}"
            if drop_first_heading and not skipped_first_heading:
                skipped_first_heading = True
                continue
            level = min(max_heading_level, base_heading_level + hashes - 1)
            if level < 1:
                level = 1
            add_numbered_heading(doc, nums, title, level)
            continue

        # Numbered list (1. ...)
        m = re.match(r"^\d+\.\s+(.*)$", s)
        if m:
            _add_bullet(doc, _strip_inline_md(m.group(1).strip()), numbered=True)
            continue

        # Bullet list (-/*)
        m = re.match(r"^[-*+]\s+(.*)$", s)
        if m:
            _add_bullet(doc, _strip_inline_md(m.group(1).strip()), numbered=False)
            continue

        # Normal paragraph
        doc.add_paragraph(_strip_inline_md(s), style="Normal")

    if pending_table:
        flush_table()
