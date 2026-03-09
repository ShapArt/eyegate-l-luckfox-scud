from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import sys
import textwrap
import urllib.request
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt
from PIL import Image

# Official Luckfox wiki materials (downloaded as "reference for prototype", not project-owned KD).
LUCKFOX_PINOUT_JPG_URL = "https://wiki.luckfox.com/assets/images/Luckfox-Pico-Ultra-details-inter-e55ff8f5f76d54e25629705935d42765.jpg"
LUCKFOX_ULTRA_W_PDF_URL = "https://wiki.luckfox.com/assets/files/Luckfox-Pico-Ultra-W-5fb258ed24c23939a00297232d84e214.pdf"


OUT_DIRNAME = "КП_2025_EyeGate_Mantrap_Шаповалов_ИУ8-74"
DOCSET_DIRNAME = "КП_2025_EyeGate_Mantrap_DOCS"


@dataclass(frozen=True)
class RegistryRow:
    n: int
    title: str
    basis: str
    relpath: str
    on_server: str  # "ДА"/"НЕТ"
    note: str = ""


def _find_repo_root(start: Path) -> Path:
    cur = start.resolve()
    for candidate in (cur, *cur.parents):
        if (candidate / ".git").exists():
            return candidate
    return start.resolve()


def _ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _snapshot_to_archive(repo_root: Path, paths: list[Path], *, label: str) -> Path:
    snap_root = repo_root / "archive" / f"{label}_{_timestamp()}"
    _ensure_dir(snap_root)
    for p in paths:
        if not p.exists():
            continue
        dst = snap_root / p.name
        if p.is_dir():
            shutil.copytree(p, dst, dirs_exist_ok=True)
        else:
            shutil.copy2(p, dst)
    return snap_root


def _extract_drive_download_zips(repo_root: Path) -> list[Path]:
    inbox = repo_root / "_inbox"
    _ensure_dir(inbox)
    extracted: list[Path] = []
    zips = sorted(repo_root.glob("drive-download-*.zip"))
    if not zips:
        # Repo cleanup may move raw inputs into archive/inputs/.
        zips = sorted((repo_root / "archive" / "inputs").glob("drive-download-*.zip"))
    for z in zips:
        dest = inbox / z.stem
        if dest.exists() and any(dest.rglob("*")):
            extracted.append(dest)
            continue
        _ensure_dir(dest)
        with zipfile.ZipFile(z, "r") as zf:
            zf.extractall(dest)
        extracted.append(dest)
    return extracted


def _field_run(paragraph, instruction: str) -> None:
    """
    Insert a Word field (complex form).

    NOTE: Word fields should be represented as multiple runs (begin/instr/separate/result/end).
    A single-run field is flaky and often doesn't update/render correctly.
    """
    r_begin = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    r_begin._r.append(begin)

    r_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    r_instr._r.append(instr)

    r_sep = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    r_sep._r.append(separate)

    paragraph.add_run(" ")

    r_end = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r_end._r.append(end)


def _enable_update_fields_on_open(doc: Document) -> None:
    """Ask Word to update TOC/PAGE fields automatically on open."""
    try:
        settings = doc.settings.element
        upd = settings.find(qn("w:updateFields"))
        if upd is None:
            upd = OxmlElement("w:updateFields")
            settings.append(upd)
        upd.set(qn("w:val"), "true")
    except Exception:
        pass


def _configure_page(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(10)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(14)


def _set_footer_page_numbers(doc: Document) -> None:
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True

    # Headers empty.
    for hdr in (sec.header, sec.first_page_header):
        hdr.is_linked_to_previous = False
        if not hdr.paragraphs:
            hdr.add_paragraph("")
        for p in hdr.paragraphs:
            p.text = ""

    # Primary footer: centered PAGE field.
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field_run(p, "PAGE")

    # First page footer: empty.
    fp = sec.first_page_footer
    fp.is_linked_to_previous = False
    if not fp.paragraphs:
        fp.add_paragraph("")
    for p in fp.paragraphs:
        p.text = ""


def _title_page(doc: Document, meta: dict[str, Any], doc_title: str) -> None:
    # Title page layout based on LR3_ASVT_Lukyanov_IU8-64.docx (к афедральный шаблон),
    # but strictly without floating drawings (wp:anchor) to pass norm-control.
    repo_root = _find_repo_root(Path(__file__).resolve())
    lr3 = repo_root / "LR3_ASVT_Lukyanov_IU8-64.docx"

    def _initials(full_name: str) -> str:
        parts = [p for p in (full_name or "").split() if p.strip()]
        if not parts:
            return ""
        fam = parts[0]
        inits = "".join((p[0].upper() + ".") for p in parts[1:3] if p)
        return (fam + (" " + inits if inits else "")).strip()

    def _no_proof_run(run) -> None:
        # Avoid red spellcheck underlines on title pages (print output is unaffected, but it looks cleaner).
        try:
            rpr = run._r.get_or_add_rPr()
            if rpr.find(qn("w:noProof")) is None:
                rpr.append(OxmlElement("w:noProof"))
        except Exception:
            pass

    def _set_table_borders_none(table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        for b in list(tblPr.findall(qn("w:tblBorders"))):
            tblPr.remove(b)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
        tblPr.append(borders)

    def _set_table_fixed_layout(table) -> None:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        layout = tblPr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tblPr.append(layout)
        layout.set(qn("w:type"), "fixed")

    def _set_cell_margins_zero(cell) -> None:
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = tcPr.find(qn("w:tcMar"))
            if tcMar is None:
                tcMar = OxmlElement("w:tcMar")
                tcPr.append(tcMar)
            for edge in ("top", "left", "bottom", "right"):
                el = tcMar.find(qn(f"w:{edge}"))
                if el is None:
                    el = OxmlElement(f"w:{edge}")
                    tcMar.append(el)
                el.set(qn("w:w"), "0")
                el.set(qn("w:type"), "dxa")
        except Exception:
            pass

    def _lr3_logo_bytes() -> bytes | None:
        if not lr3.exists():
            return None
        try:
            with zipfile.ZipFile(lr3, "r") as z:
                for cand in (
                    "word/media/image1.jpeg",
                    "word/media/image1.jpg",
                    "word/media/image1.png",
                ):
                    if cand in z.namelist():
                        return z.read(cand)
                for name in z.namelist():
                    if not name.startswith("word/media/"):
                        continue
                    low = name.lower()
                    if low.endswith((".jpeg", ".jpg", ".png")):
                        return z.read(name)
        except Exception:
            return None
        return None

    uni = meta.get("university") or {}
    edu = meta.get("education") or {}
    people = meta.get("people") or {}
    proj = meta.get("project") or {}

    city = str(uni.get("city") or "Москва")
    year = str(uni.get("year") or "")

    t1 = doc.add_table(rows=1, cols=2)
    t1.autofit = False
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders_none(t1)
    _set_table_fixed_layout(t1)
    w_logo = Cm(2.44)
    w_text = Cm(14.56)
    for row in t1.rows:
        row.cells[0].width = w_logo
        row.cells[1].width = w_text
        for c in row.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # Ensure tblGrid columns match our intended widths (prevents layout drift).
    try:
        t1.columns[0].width = w_logo
        t1.columns[1].width = w_text
    except Exception:
        pass

    logo_cell = t1.cell(0, 0)
    logo_cell.text = ""
    lp = logo_cell.paragraphs[0]
    lp.paragraph_format.first_line_indent = None
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)
    lp.paragraph_format.line_spacing = 1.0
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_cell_margins_zero(logo_cell)
    try:
        from io import BytesIO

        logo = _lr3_logo_bytes()
        if logo:
            # Ensure no stray whitespace/run before the picture.
            lp.text = ""
            lp.add_run().add_picture(BytesIO(logo), width=Cm(2.10))
    except Exception:
        pass

    header_lines = [
        "Министерство науки и высшего образования Российской Федерации",
        "Федеральное государственное бюджетное образовательное учреждение ",
        "высшего образования",
        "«Московский государственный технический университет",
        "имени Н.Э. Баумана",
        "(национальный исследовательский университет)»",
        "(МГТУ им. Н.Э. Баумана)",
    ]
    hc = t1.cell(0, 1)
    hc.text = ""
    _set_cell_margins_zero(hc)
    for i, line in enumerate(header_lines):
        p = hc.paragraphs[0] if i == 0 else hc.add_paragraph()
        p.text = (line or "").rstrip()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
            _no_proof_run(r)

    doc.add_paragraph("")
    doc.add_paragraph("")

    t2 = doc.add_table(rows=2, cols=2)
    t2.autofit = False
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders_none(t2)
    _set_table_fixed_layout(t2)
    # Fit exactly into text area width: 17.0 cm (A4 minus margins 3cm+1cm).
    w_lbl = Cm(4.0)
    w_val = Cm(13.0)
    for row in t2.rows:
        row.cells[0].width = w_lbl
        row.cells[1].width = w_val
        for c in row.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    try:
        t2.columns[0].width = w_lbl
        t2.columns[1].width = w_val
    except Exception:
        pass

    t2.cell(0, 0).text = "Факультет"
    t2.cell(0, 1).text = "«Информатика и системы управления» (ИУ)"
    t2.cell(1, 0).text = "Кафедра"
    t2.cell(1, 1).text = "«Информационная безопасность» (ИУ8)"
    for row in t2.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    _no_proof_run(r)
    # Prevent ugly word wrapping in the left label column (e.g., "Факульт\nет").
    for r in range(2):
        try:
            tcPr = t2.cell(r, 0)._tc.get_or_add_tcPr()
            no_wrap = OxmlElement("w:noWrap")
            tcPr.append(no_wrap)
        except Exception:
            pass

    doc.add_paragraph("")

    def _title_line(text: str, *, bold: bool) -> None:
        p = doc.add_paragraph("", style="Normal")
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        _no_proof_run(run)

    _title_line("КУРСОВОЙ ПРОЕКТ", bold=True)
    title2 = str(doc_title).upper()
    if len(title2) > 50:
        p = doc.add_paragraph("", style="Normal")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title2)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        _no_proof_run(run)
    else:
        _title_line(title2, bold=True)
    theme = str(proj.get("theme_ru") or proj.get("name_ru") or "").strip()
    if theme:
        theme = theme.rstrip("/").strip()
        p = doc.add_paragraph("", style="Normal")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"«{theme}»")
        run.bold = False
        run.font.name = "Times New Roman"
        if len(theme) > 110:
            run.font.size = Pt(11)
        elif len(theme) > 80:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(14)
        _no_proof_run(run)

    # Keep it adaptive: long titles may wrap, so reduce empty paragraphs to prevent spill to page 2.
    blank_n = 8
    if len(str(doc_title)) > 70 or len(theme) > 120:
        blank_n = 6
    if len(str(doc_title)) > 100 or len(theme) > 160:
        blank_n = 4
    for _ in range(blank_n):
        p = doc.add_paragraph("", style="Normal")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    t3 = doc.add_table(rows=2, cols=2)
    t3.autofit = False
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders_none(t3)
    _set_table_fixed_layout(t3)
    # Fit exactly into text area width (17.0 cm).
    w_left = Cm(11.8)
    w_right = Cm(5.2)
    for row in t3.rows:
        row.cells[0].width = w_left
        row.cells[1].width = w_right
        for c in row.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    try:
        t3.columns[0].width = w_left
        t3.columns[1].width = w_right
    except Exception:
        pass

    supervisor = str(people.get("supervisor") or "").strip()
    student_full = str(people.get("student") or "").strip()
    student = _initials(student_full) or student_full

    t3.cell(0, 0).text = f"Руководитель:\n{supervisor}".rstrip()
    # Title page: student name only (group is not shown per supervisor request).
    t3.cell(1, 0).text = f"Студент:\n{student}".rstrip()
    for row in t3.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                for r in p.runs:
                    _no_proof_run(r)

    def _sign_cell(cell) -> None:
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.text = "____________________"
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.first_line_indent = None
        p2 = cell.add_paragraph("(подпись, дата)")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = None

    _sign_cell(t3.cell(0, 1))
    _sign_cell(t3.cell(1, 1))

    # City/year goes into first-page footer (no page number on title page, but the title page looks complete).
    if year:
        fp = doc.sections[0].first_page_footer
        fp.is_linked_to_previous = False
        if not fp.paragraphs:
            fp.add_paragraph("")
        p = fp.paragraphs[0]
        p.text = f"{city}, {year}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            _no_proof_run(r)

    doc.add_page_break()


def _download(url: str, dest: Path) -> None:
    _ensure_dir(dest.parent)
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=60) as r:  # noqa: S310
        data = r.read()
    dest.write_bytes(data)


def _download_luckfox_refs(out_dir: Path) -> dict[str, str]:
    """
    Returns dict with keys:
    - pdf_path
    - pinout_png_path
    - pinout_src_url
    - pdf_src_url
    """
    refs_dir = out_dir / "07_Графика_и_презентация" / "Luckfox_официально"
    _ensure_dir(refs_dir)

    pdf_path = refs_dir / "Luckfox_Pico_Ultra_W_official.pdf"
    jpg_path = refs_dir / "Luckfox_Pico_Ultra_W_pinout_official.jpg"
    png_path = refs_dir / "Luckfox_Pico_Ultra_W_pinout_official.png"

    if not pdf_path.exists():
        _download(LUCKFOX_ULTRA_W_PDF_URL, pdf_path)
    if not jpg_path.exists():
        _download(LUCKFOX_PINOUT_JPG_URL, jpg_path)
    if not png_path.exists():
        with Image.open(jpg_path) as im:
            im.convert("RGB").save(png_path)

    readme = refs_dir / "README_SOURCES.md"
    readme.write_text(
        textwrap.dedent(
            f"""\
            # Luckfox Pico Ultra W (официальные материалы)

            Сохранено как справочный материал прототипа (не КД серийного изделия проекта).

            - PDF (Ultra W): {pdf_path.name}
              Source: {LUCKFOX_ULTRA_W_PDF_URL}
            - Pinout image: {png_path.name} (конвертировано из JPG)
              Source: {LUCKFOX_PINOUT_JPG_URL}

            Downloaded: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
            """
        ).strip()
        + "\n",
        encoding="utf-8",
    )

    return {
        "pdf_path": str(pdf_path),
        "pinout_png_path": str(png_path),
        "pinout_src_url": LUCKFOX_PINOUT_JPG_URL,
        "pdf_src_url": LUCKFOX_ULTRA_W_PDF_URL,
    }


def _next_caption_number(doc: Document, caption_word: str) -> int:
    import re

    rx = re.compile(rf"^\s*{caption_word}\s+(\d+)\s*[-–—]\s*", flags=re.IGNORECASE)
    mx = 0
    for p in doc.paragraphs:
        m = rx.match((p.text or "").strip())
        if not m:
            continue
        try:
            mx = max(mx, int(m.group(1)))
        except Exception:
            continue
    return mx + 1


def _append_luckfox_pinout_figure(
    docx_path: Path, pinout_png: Path, *, source_url: str
) -> None:
    doc = Document(str(docx_path))
    fig_n = _next_caption_number(doc, "Рисунок")

    # Simple unnumbered heading to avoid breaking manual numbering in existing docs.
    h = doc.add_paragraph(
        "Справочный материал прототипа: Luckfox Pico Ultra W", style="Heading 1"
    )
    h.paragraph_format.first_line_indent = None

    p_img = doc.add_paragraph("", style="Normal")
    p_img.paragraph_format.first_line_indent = None
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(pinout_png), width=Cm(16.5))

    cap = doc.add_paragraph(
        f"Рисунок {fig_n} – Пинаут и разъёмы Luckfox Pico Ultra W (официальная схема выводов)",
        style="Normal",
    )
    cap.paragraph_format.first_line_indent = None
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    src = doc.add_paragraph(f"Источник: {source_url}", style="Normal")
    src.paragraph_format.first_line_indent = None

    doc.save(str(docx_path))


def _write_inventory(repo_root: Path, out_dir: Path) -> Path:
    exts = {
        ".docx",
        ".pdf",
        ".pptx",
        ".drawio",
        ".vsdx",
        ".png",
        ".svg",
        ".jpg",
        ".jpeg",
    }
    skip_dirs = {
        ".git",
        ".venv",
        "__pycache__",
        "node_modules",
        "archive",
        "_inbox",
        OUT_DIRNAME,
    }
    items: list[dict[str, Any]] = []
    for p in repo_root.rglob("*"):
        if not p.is_file():
            continue
        if p.suffix.lower() not in exts:
            continue
        rel_parts = p.relative_to(repo_root).parts
        if rel_parts and rel_parts[0] in skip_dirs:
            continue
        items.append(
            {
                "path": str(p.relative_to(repo_root)),
                "size": p.stat().st_size,
            }
        )
    inv_path = out_dir / "_meta" / "inventory.json"
    _ensure_dir(inv_path.parent)
    inv_path.write_text(
        json.dumps(items, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return inv_path


def _copy_file(src: Path, dst: Path) -> None:
    _ensure_dir(dst.parent)
    shutil.copy2(src, dst)


def _copy_tree(
    src: Path, dst: Path, *, ignore: shutil.IgnorePattern | None = None
) -> None:
    _ensure_dir(dst.parent)
    shutil.copytree(src, dst, dirs_exist_ok=True, ignore=ignore)


def _copy_repo_sources(repo_root: Path, out_dir: Path) -> None:
    dst_root = out_dir / "Приложение_Исходники_и_сборка" / "repo"
    _ensure_dir(dst_root)

    # Keep the "model package" lightweight/rebuildable:
    # - exclude installed deps/caches and generated artifacts (node_modules, dist, __pycache__, etc.)
    # - exclude large bundled binaries (ffmpeg/mediamtx/adb) since they are not required for code review;
    #   project scripts support using system-installed tools as well.
    #
    # If you need a fully offline package later, remove these ignores and rebuild.
    ignore = shutil.ignore_patterns(
        ".git",
        ".venv",
        "__pycache__",
        ".pytest_cache",
        "node_modules",
        "dist",
        "*.pyc",
        "*.pyo",
        "*.log",
        "*.tmp",
        "*.swp",
        "*.DS_Store",
        "*:Zone.Identifier",
        "archive",
        "_inbox",
        OUT_DIRNAME,
        DOCSET_DIRNAME,
        "ffmpeg",
        "mediamtx",
        "adb",
    )

    include_dirs = [
        "server",
        "gate",
        "vision",
        "camera_ingest",
        "hw",
        "db",
        "auth",
        "policy",
        "web",
        "scripts",
        "deploy",
        "luckfox",
        "models",
        "docs",
        "tools",
        "data",
        "tests",
    ]
    include_files = [
        "README.md",
        "README_QUICKSTART.txt",
        "AGENTS.md",
        ".env.example",
        "requirements.txt",
        "requirements-dev.txt",
        "Makefile",
        "pytest.ini",
    ]

    for d in include_dirs:
        src = repo_root / d
        if src.is_dir():
            _copy_tree(src, dst_root / d, ignore=ignore)
    for f in include_files:
        src = repo_root / f
        if src.is_file():
            _copy_file(src, dst_root / f)

    # Keep archives in the main repo; just link where they are.
    (dst_root / "README_ARCHIVES.md").write_text(
        textwrap.dedent(
            """\
            # Large archives

            Большие исходные ZIP/7Z (drive-download-*.zip, eyegate-mantrap.7z и т.п.) вынесены в `archive/inputs/`,
            чтобы не раздувать папку сдачи. При необходимости их можно приложить отдельным архивом.
            """
        ).strip()
        + "\n",
        encoding="utf-8",
    )


def _make_norm_table(
    doc: Document, headers: list[str], rows: list[list[str]], col_widths_cm: list[float]
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    # Widths and formatting
    for i, w in enumerate(col_widths_cm):
        for r in table.rows:
            if i < len(r.cells):
                r.cells[i].width = Cm(w)
                r.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, r in enumerate(table.rows):
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run.bold = bool(r_idx == 0)


def _generate_registry_doc(
    out_dir: Path, meta: dict[str, Any], rows: list[RegistryRow]
) -> Path:
    doc = Document()
    _configure_page(doc)
    _set_footer_page_numbers(doc)
    _enable_update_fields_on_open(doc)
    _title_page(doc, meta, "Реестр передачи на сервер кафедры")

    p = doc.add_paragraph(
        "Реестр передачи на сервер кафедры (комплект КП_2025).", style="Normal"
    )
    p.paragraph_format.first_line_indent = None

    headers = [
        "№",
        "Документ/файл",
        "ГОСТ/основание",
        "Где лежит",
        "На сервер кафедры",
        "Примечание",
    ]
    table_rows = [[r.n, r.title, r.basis, r.relpath, r.on_server, r.note] for r in rows]
    col_widths = [1.0, 6.0, 3.5, 4.0, 1.5, 1.0]  # sum ~= 17.0 cm
    _make_norm_table(doc, headers, table_rows, col_widths_cm=col_widths)

    out_path = out_dir / "00_Реестр_передачи_на_сервер.docx"
    doc.save(str(out_path))
    return out_path


def _generate_model_test_results_doc(
    out_dir: Path, meta: dict[str, Any], *, pytest_log_relpath: str, pytest_rc: int
) -> Path:
    doc = Document()
    _configure_page(doc)
    _set_footer_page_numbers(doc)
    _enable_update_fields_on_open(doc)
    _title_page(doc, meta, "Испытания модели: результаты и выводы")

    p = doc.add_paragraph(
        "Результаты получены в ходе запуска контрольных примеров и автоматизированных тестов.",
        style="Normal",
    )
    p.paragraph_format.first_line_indent = None

    doc.add_paragraph(
        "1 Общие сведения", style="Heading 1"
    ).paragraph_format.first_line_indent = None
    doc.add_paragraph(
        f"Команда: pytest -q. Код возврата: {pytest_rc}. Лог: {pytest_log_relpath}. Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}.",
        style="Normal",
    ).paragraph_format.first_line_indent = None

    doc.add_paragraph(
        "2 Итоговая таблица", style="Heading 1"
    ).paragraph_format.first_line_indent = None
    headers = [
        "№",
        "Проверка",
        "Ожидаемый результат",
        "Фактический результат",
        "Статус",
        "Примечание",
    ]
    status = "УСПЕХ" if pytest_rc == 0 else "ОШИБКА"
    rows = [
        [
            "1",
            "Автотесты backend (pytest)",
            "Все тесты проходят",
            f"См. лог {pytest_log_relpath}",
            status,
            "",
        ],
    ]
    _make_norm_table(doc, headers, rows, col_widths_cm=[1.0, 4.0, 3.5, 5.0, 2.0, 1.5])

    doc.add_paragraph(
        "3 Выводы", style="Heading 1"
    ).paragraph_format.first_line_indent = None
    if pytest_rc == 0:
        doc.add_paragraph(
            "По результатам испытаний подтверждена работоспособность модели в рамках выполненных проверок.",
            style="Normal",
        ).paragraph_format.first_line_indent = None
    else:
        doc.add_paragraph(
            "По результатам испытаний выявлены ошибки (см. лог). Требуется устранение замечаний и повторный прогон проверок.",
            style="Normal",
        ).paragraph_format.first_line_indent = None

    out_path = (
        out_dir / "02_Моделирование" / "04_Испытания_модели_результаты_и_выводы.docx"
    )
    _ensure_dir(out_path.parent)
    doc.save(str(out_path))
    return out_path


def _run_pytest(repo_root: Path, out_dir: Path) -> tuple[int, Path]:
    log_dir = out_dir / "02_Моделирование" / "logs"
    _ensure_dir(log_dir)
    log_path = log_dir / "pytest.txt"
    cmd = [sys.executable, "-m", "pytest", "-q"]
    cp = subprocess.run(
        cmd, cwd=str(repo_root), capture_output=True, text=True
    )  # noqa: S603
    log_path.write_text(
        (cp.stdout or "") + "\n" + (cp.stderr or ""), encoding="utf-8", errors="replace"
    )
    return cp.returncode, log_path


def _load_meta_from_docset(docset_root: Path) -> dict[str, Any]:
    meta_path = docset_root / "_meta" / "project_meta.yaml"
    if not meta_path.exists():
        return {}
    return yaml.safe_load(meta_path.read_text(encoding="utf-8")) or {}


def _copy_docset_into_structure(repo_root: Path, out_dir: Path) -> list[RegistryRow]:
    docset_root = repo_root / DOCSET_DIRNAME
    rows: list[RegistryRow] = []

    # Create required structure.
    for d in (
        "01_ТЗ_и_титулы",
        "02_Моделирование",
        "03_КД_ЕСКД",
        "04_ПД_ЕСПД_ГОСТ19",
        "05_АСУД_ГОСТ34",
        "06_Эксплуатационная_документация",
        "07_Графика_и_презентация",
        "Приложение_Исходники_и_сборка",
    ):
        _ensure_dir(out_dir / d)

    mapping: list[tuple[str, str, str, str]] = [
        # (src filename, dst relative path, basis, note)
        (
            "01_УЧЕБНОЕ_ЗАДАНИЕ.docx",
            "01_ТЗ_и_титулы/01_Учебное_задание.docx",
            "Учебное задание (кафедра)",
            "",
        ),
        (
            "02_ТЗ_АС_ГОСТ34.docx",
            "01_ТЗ_и_титулы/02_ТЗ_АС_ГОСТ34.602-89.docx",
            "ГОСТ 34.602-89",
            "",
        ),
        (
            "03_ТЗ_ПРОГРАММА_ГОСТ19_201.docx",
            "01_ТЗ_и_титулы/03_ТЗ_на_программу_ГОСТ19.201-78.docx",
            "ГОСТ 19.201-78",
            "",
        ),
        (
            "04_ТИТУЛЬНЫЕ_ЛИСТЫ_ВИЗА.docx",
            "01_ТЗ_и_титулы/04_Титульные_листы_и_виза.docx",
            "ГОСТ 2.105-95",
            "",
        ),
        (
            "05_МОДЕЛИРОВАНИЕ_ОПИСАНИЕ_МОДЕЛИ.docx",
            "02_Моделирование/01_Описание_модели.docx",
            "Материалы КП (модель)",
            "",
        ),
        (
            "06_МОДЕЛИРОВАНИЕ_КОНТРОЛЬНЫЕ_ПРИМЕРЫ.docx",
            "02_Моделирование/02_Контрольные_примеры.docx",
            "Материалы КП (модель)",
            "",
        ),
        (
            "07_ПРОГРАММА_И_МЕТОДИКА_ИСПЫТАНИЙ_МОДЕЛИ.docx",
            "02_Моделирование/03_ПМИ_модели.docx",
            "ПМИ модели (учебная)",
            "",
        ),
        (
            "21_ПСИ_ПРОГРАММА_И_МЕТОДИКА_ИСПЫТАНИЙ.docx",
            "02_Моделирование/05_ПСИ_программа_и_методика_испытаний.docx",
            "ПСИ (учебная)",
            "",
        ),
        (
            "22_ПСИ_ПРОТОКОЛ_РЕЗУЛЬТАТОВ_ШАБЛОН.docx",
            "02_Моделирование/06_ПСИ_протокол_результатов_шаблон.docx",
            "ПСИ (учебная)",
            "Шаблон протокола; фактические результаты см. в 04_Испытания_модели_...",
        ),
        (
            "00_РПЗ_РАСЧЕТНО_ПОЯСНИТЕЛЬНАЯ_ЗАПИСКА.docx",
            "05_АСУД_ГОСТ34/РПЗ.docx",
            "РД 50-34.698-90, ГОСТ 2.105-95",
            "",
        ),
        (
            "08_СХЕМА_СТРУКТУРНАЯ_ФУНКЦИОНАЛЬНАЯ.docx",
            "03_КД_ЕСКД/08_Схема_структурная_функциональная.docx",
            "ЕСКД (учебная)",
            "",
        ),
        (
            "09_СХЕМА_ЭЛЕКТРИЧЕСКАЯ_ПРИНЦИПИАЛЬНАЯ_Э3.docx",
            "03_КД_ЕСКД/09_Схема_электрическая_Э3.docx",
            "ЕСКД (учебная)",
            "",
        ),
        (
            "10_ПЕРЕЧЕНЬ_ЭЛЕМЕНТОВ_ПЭ.docx",
            "03_КД_ЕСКД/10_Перечень_элементов_ПЭ.docx",
            "ЕСКД (учебная)",
            "",
        ),
        (
            "11_ЧЕРТЕЖ_ПП_ЦЕНТРАЛЬНЫЙ_КОНТРОЛЛЕР.docx",
            "03_КД_ЕСКД/11_Чертеж_ПП_центральный_контроллер_эскиз.docx",
            "ЕСКД (учебная)",
            "",
        ),
        (
            "12_СБОРОЧНЫЙ_ЧЕРТЕЖ_ПП.docx",
            "03_КД_ЕСКД/12_Сборочный_чертеж_ПП.docx",
            "ЕСКД (учебная)",
            "",
        ),
        (
            "13_СПЕЦИФИКАЦИЯ.docx",
            "03_КД_ЕСКД/13_Спецификация.docx",
            "ЕСКД (учебная)",
            "",
        ),
        (
            "14_СХЕМЫ_АЛГОРИТМОВ_ГОСТ19_701.docx",
            "04_ПД_ЕСПД_ГОСТ19/14_Схемы_алгоритмов_ГОСТ19.701-90.docx",
            "ГОСТ 19.701-90",
            "",
        ),
        (
            "15_СТРУКТУРНАЯ_СХЕМА_ПО.docx",
            "04_ПД_ЕСПД_ГОСТ19/15_Структурная_схема_ПО.docx",
            "ЕСПД (учебная)",
            "",
        ),
        (
            "16_ИСХОДНЫЕ_ТЕКСТЫ_ПРОГРАММЫ_ГОСТ19_401.docx",
            "04_ПД_ЕСПД_ГОСТ19/16_Исходные_тексты_программы_ГОСТ19.401.docx",
            "ГОСТ 19.401",
            "",
        ),
        (
            "17_ОПИСАНИЕ_ПРОГРАММЫ_ГОСТ19_402.docx",
            "04_ПД_ЕСПД_ГОСТ19/17_Описание_программы_ГОСТ19.402.docx",
            "ГОСТ 19.402",
            "",
        ),
        (
            "18_РУКОВОДСТВО_ПОЛЬЗОВАТЕЛЯ_ОПЕРАТОРА.docx",
            "06_Эксплуатационная_документация/18_Руководство_оператора.docx",
            "Эксплуатационная документация",
            "",
        ),
        (
            "19_РУКОВОДСТВО_СИСАДМИНА.docx",
            "06_Эксплуатационная_документация/19_Руководство_администратора.docx",
            "Эксплуатационная документация",
            "",
        ),
        (
            "20_РУКОВОДСТВО_ПРОГРАММИСТА.docx",
            "06_Эксплуатационная_документация/20_Руководство_программиста.docx",
            "Эксплуатационная документация",
            "",
        ),
        (
            "23_ПРИЛОЖЕНИЯ_СБОРКА_И_ЗАПУСК.docx",
            "Приложение_Исходники_и_сборка/01_Сборка_и_запуск.docx",
            "Приложение к РПЗ",
            "",
        ),
    ]

    n = 1
    for src_name, dst_rel, basis, note in mapping:
        src = docset_root / src_name
        if not src.exists():
            continue
        dst = out_dir / dst_rel
        _copy_file(src, dst)
        rows.append(
            RegistryRow(
                n=n,
                title=dst.name,
                basis=basis,
                relpath=dst_rel,
                on_server="ДА",
                note=note,
            )
        )
        n += 1

    # Copy traceability matrix and QA checklist
    for extra in (
        "TRACEABILITY_MATRIX.csv",
        "TRACEABILITY_MATRIX.xlsx",
        "QA_CHECKLIST.md",
    ):
        src = docset_root / extra
        if src.exists():
            dst_rel = f"05_АСУД_ГОСТ34/{extra}"
            _copy_file(src, out_dir / dst_rel)
            rows.append(
                RegistryRow(
                    n=n,
                    title=extra,
                    basis="Материалы КП (трассировка/контроль)",
                    relpath=dst_rel,
                    on_server="ДА",
                )
            )
            n += 1

    # Posters + presentation
    posters_src = docset_root / "24_ПЛАКАТЫ_И_СХЕМЫ"
    if posters_src.is_dir():
        posters_dst_rel = "07_Графика_и_презентация/Плакаты_и_схемы"
        _copy_tree(posters_src, out_dir / posters_dst_rel)
        rows.append(
            RegistryRow(
                n=n,
                title="Плакаты_и_схемы (PNG/PDF)",
                basis="Материалы КП (графика)",
                relpath=posters_dst_rel,
                on_server="ДА",
            )
        )
        n += 1

    pres_src = docset_root / "25_ПРЕЗЕНТАЦИЯ_КП.pptx"
    if pres_src.exists():
        pres_dst_rel = "07_Графика_и_презентация/Презентация_КП.pptx"
        _copy_file(pres_src, out_dir / pres_dst_rel)
        rows.append(
            RegistryRow(
                n=n,
                title="Презентация_КП.pptx",
                basis="Материалы КП (презентация)",
                relpath=pres_dst_rel,
                on_server="ДА",
            )
        )
        n += 1

    # Diagrams (sources + exports)
    schemes_src = docset_root / "schemes"
    if schemes_src.is_dir():
        schemes_dst_rel = "03_КД_ЕСКД/schemes"
        _copy_tree(schemes_src, out_dir / schemes_dst_rel)
        rows.append(
            RegistryRow(
                n=n,
                title="schemes (drawio/svg/png)",
                basis="ЕСКД/ЕСПД (схемы, учебная)",
                relpath=schemes_dst_rel,
                on_server="ДА",
            )
        )
        n += 1

    # Docset appendix/model files as build helpers
    for dname in ("appendix", "model_files", "scans"):
        src = docset_root / dname
        if not src.is_dir():
            continue
        dst_rel = f"Приложение_Исходники_и_сборка/{dname}"
        _copy_tree(src, out_dir / dst_rel)
        rows.append(
            RegistryRow(
                n=n,
                title=dname,
                basis="Материалы КП (приложения)",
                relpath=dst_rel,
                on_server="ДА",
            )
        )
        n += 1

    return rows


def _write_missing_inputs(out_dir: Path, meta: dict[str, Any]) -> Path:
    missing: list[str] = []
    sup = ((meta.get("people") or {}).get("supervisor") or "").strip()
    if not sup:
        missing.append(
            "ФИО руководителя/консультанта (как писать в титульных листах) + должность/степень (если требуется)."
        )
    else:
        missing.append(
            f"Если кафедра требует полное ФИО руководителя + должность/степень: уточнить (сейчас в документах указано: «{sup}»)."
        )

    missing.append(
        "Если кафедра требует номенклатуру вида МГТУ.XXXXXX.XXX: сообщить обозначения для ТЗ/РПЗ/ПМИ/руководств; "
        "иначе в документах используется код проекта (КП_2025_EyeGate_Mantrap)."
    )

    p = out_dir / "MISSING_INPUTS.md"
    p.write_text(
        "# Missing inputs\n\n" + "\n".join(f"- {x}" for x in missing) + "\n",
        encoding="utf-8",
    )
    return p


def _zip_subset(src_root: Path, zip_path: Path, *, include_relpath) -> None:
    """
    Create a ZIP with a subset of files from src_root.

    include_relpath(rel: Path) -> bool
      where rel is relative to src_root.
    """
    _ensure_dir(zip_path.parent)
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(
        zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as zf:
        files = [p for p in src_root.rglob("*") if p.is_file()]
        for p in sorted(files):
            rel = p.relative_to(src_root)
            if not include_relpath(rel):
                continue
            arc = (Path(src_root.name) / rel).as_posix()
            zf.write(p, arcname=arc)


def _make_dist_zips(repo_root: Path, out_dir: Path) -> tuple[Path, Path]:
    """
    Produce the two teacher-friendly archives:
    - Документация: only docs/graphics (no full repo sources).
    - Модель: appendix with sources/build scripts.
    """
    dist_dir = repo_root / "dist"
    _ensure_dir(dist_dir)

    doc_zip = dist_dir / f"{out_dir.name}_Документация.zip"
    model_zip = dist_dir / f"{out_dir.name}_Модель.zip"

    appendix_dirname = "Приложение_Исходники_и_сборка"
    _zip_subset(
        out_dir,
        doc_zip,
        include_relpath=lambda rel: not (
            rel.parts and rel.parts[0] == appendix_dirname
        ),
    )
    _zip_subset(
        out_dir,
        model_zip,
        include_relpath=lambda rel: bool(rel.parts)
        and rel.parts[0] == appendix_dirname,
    )
    return doc_zip, model_zip


def main(argv: list[str]) -> int:
    ap = argparse.ArgumentParser(
        description="Build KP_2025 submission bundle (Word-first, norm-control oriented)."
    )
    ap.add_argument(
        "--skip-unzip",
        action="store_true",
        help="Do not extract drive-download-*.zip into _inbox/",
    )
    ap.add_argument(
        "--skip-docset", action="store_true", help="Do not run tools/docs/build_docs.py"
    )
    ap.add_argument(
        "--skip-pytest",
        action="store_true",
        help="Do not run pytest -q (results doc will be created without log)",
    )
    ap.add_argument(
        "--skip-repo-copy",
        action="store_true",
        help="Do not copy repo sources into appendix",
    )
    ap.add_argument(
        "--skip-style-check",
        action="store_true",
        help="Do not run tools/docx_style_check.py",
    )
    args = ap.parse_args(argv)

    repo_root = _find_repo_root(Path(__file__).resolve())

    # Ensure archives folder exists.
    _ensure_dir(repo_root / "archive")

    if not args.skip_unzip:
        _extract_drive_download_zips(repo_root)

    # Snapshot previous output (if any) before overwriting.
    out_dir = repo_root / OUT_DIRNAME
    legacy_out_dirs = [
        repo_root / "КП_2025_EyeGate_Mantrap_Яшев_ИУ8-74",
    ]
    existing_to_snapshot: list[Path] = []
    if out_dir.exists():
        existing_to_snapshot.append(out_dir)
    for p in legacy_out_dirs:
        if p.exists():
            existing_to_snapshot.append(p)
    # Also snapshot generated docset + previous FINAL zip if present.
    docset_root = repo_root / DOCSET_DIRNAME
    if docset_root.exists():
        existing_to_snapshot.append(docset_root)
    final_zip = repo_root / "КП_2025_EyeGate_Mantrap_DOCS_FINAL.zip"
    if final_zip.exists():
        existing_to_snapshot.append(final_zip)
    if existing_to_snapshot:
        _snapshot_to_archive(repo_root, existing_to_snapshot, label="kp2025_snapshot")

    # Move legacy outputs out of the repo root (keep copies already snapshotted).
    for p in legacy_out_dirs:
        if not p.exists():
            continue
        dst = repo_root / "archive" / f"{p.name}_moved_{_timestamp()}"
        shutil.move(str(p), str(dst))
        (dst / "README_MOVED.md").write_text(
            "Перенесено в archive/ автоматически (старый комплект сдачи, чтобы не было дублей в корне репозитория).\n",
            encoding="utf-8",
        )

    # (Re)build docset (source for most DOCX).
    if not args.skip_docset:
        subprocess.run(
            [sys.executable, "tools/docs/build_docs.py", "--refresh-facts", "--no-zip"],
            cwd=str(repo_root),
            check=True,
        )  # noqa: S603

    # Recreate output directory.
    if out_dir.exists():
        shutil.rmtree(out_dir)
    _ensure_dir(out_dir)

    meta = _load_meta_from_docset(docset_root)
    inv_path = _write_inventory(repo_root, out_dir)

    # Copy docs into required submission structure.
    registry_rows = _copy_docset_into_structure(repo_root, out_dir)

    # Download Luckfox official refs (PDF + pinout image).
    luckfox = _download_luckfox_refs(out_dir)
    registry_rows.append(
        RegistryRow(
            n=len(registry_rows) + 1,
            title="Luckfox_официально (PDF/PNG)",
            basis="Справочные материалы прототипа (официальный источник)",
            relpath="07_Графика_и_презентация/Luckfox_официально",
            on_server="ДА",
            note="Использовано как справочный материал; КД серийного изделия не разрабатывалась.",
        )
    )

    # Insert pinout figure into key Word docs (modeling + dev manual).
    pinout_png = Path(luckfox["pinout_png_path"])
    try:
        _append_luckfox_pinout_figure(
            out_dir / "02_Моделирование" / "01_Описание_модели.docx",
            pinout_png,
            source_url=luckfox["pinout_src_url"],
        )
        _append_luckfox_pinout_figure(
            out_dir
            / "06_Эксплуатационная_документация"
            / "20_Руководство_программиста.docx",
            pinout_png,
            source_url=luckfox["pinout_src_url"],
        )
    except Exception:
        # Do not fail the whole build; style report will surface if docx got corrupted.
        pass

    # Run pytest and generate "results" doc for modeling.
    pytest_rc = 999
    pytest_log_rel = ""
    if not args.skip_pytest:
        pytest_rc, pytest_log_path = _run_pytest(repo_root, out_dir)
        pytest_log_rel = str(pytest_log_path.relative_to(out_dir))
        _generate_model_test_results_doc(
            out_dir, meta, pytest_log_relpath=pytest_log_rel, pytest_rc=pytest_rc
        )
        registry_rows.append(
            RegistryRow(
                n=len(registry_rows) + 1,
                title="04_Испытания_модели_результаты_и_выводы.docx",
                basis="Материалы КП (испытания модели)",
                relpath="02_Моделирование/04_Испытания_модели_результаты_и_выводы.docx",
                on_server="ДА",
                note=f"pytest_rc={pytest_rc}",
            )
        )
        registry_rows.append(
            RegistryRow(
                n=len(registry_rows) + 1,
                title="pytest.txt",
                basis="Логи испытаний (pytest)",
                relpath=pytest_log_rel,
                on_server="ДА",
            )
        )

    # Copy repo sources (appendix).
    if not args.skip_repo_copy:
        _copy_repo_sources(repo_root, out_dir)
        registry_rows.append(
            RegistryRow(
                n=len(registry_rows) + 1,
                title="Исходники (папка repo)",
                basis="Приложение (исходники и сборка)",
                relpath="Приложение_Исходники_и_сборка/repo",
                on_server="ДА",
                note="Содержит исходники проекта и скрипты запуска.",
            )
        )

    # Generate registry doc (must be first in folder structure).
    registry_rows_sorted: list[RegistryRow] = [
        RegistryRow(
            n=1,
            title="00_Реестр_передачи_на_сервер.docx",
            basis="Требование кафедры (реестр комплекта)",
            relpath="00_Реестр_передачи_на_сервер.docx",
            on_server="ДА",
        )
    ]
    registry_rows_sorted.extend(
        [
            RegistryRow(
                n=i + 2,
                title=r.title,
                basis=r.basis,
                relpath=r.relpath,
                on_server=r.on_server,
                note=r.note,
            )
            for i, r in enumerate(registry_rows)
        ]
    )
    _generate_registry_doc(out_dir, meta, registry_rows_sorted)

    # Missing inputs list.
    _write_missing_inputs(out_dir, meta)

    # Run style check and write report into output root.
    if not args.skip_style_check:
        subprocess.run(
            [sys.executable, "tools/docx_style_check.py", "--root", str(out_dir)],
            cwd=str(repo_root),
            check=False,
        )  # noqa: S603

    # Dist ZIPs (documentation + model).
    doc_zip, model_zip = _make_dist_zips(repo_root, out_dir)

    # Final build marker
    (out_dir / "_meta" / "build_info.json").write_text(
        json.dumps(
            {
                "built_at": datetime.now().isoformat(timespec="seconds"),
                "inventory": str(inv_path.relative_to(out_dir)),
                "pytest_rc": pytest_rc,
                "dist_doc_zip": str(doc_zip.relative_to(repo_root)),
                "dist_model_zip": str(model_zip.relative_to(repo_root)),
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
    )

    print(out_dir)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
