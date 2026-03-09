from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
TARGET = (
    PROJECT_ROOT / "КП_2025_EyeGate_Mantrap" / "02_КД_ЕСКД" / "РПЗ_EyeGate_Mantrap.docx"
)
REPO_FACTS_PATH = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "repo_facts.json"
)
TEMPLATE = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")


def load_facts() -> dict:
    if REPO_FACTS_PATH.exists():
        try:
            return json.loads(REPO_FACTS_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def copy_template() -> Document:
    if TEMPLATE.exists():
        TARGET.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(TEMPLATE, TARGET)
        doc = Document(TARGET)
        clear_body(doc)
    else:
        doc = Document()
    return doc


def insert_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def fmea_rows() -> list[list[str]]:
    return [
        [
            "Камера",
            "Нет кадра / OFF",
            "Обрыв/драйвер",
            "Событие ALARM, deny проход",
            "Смена кабеля/драйвера; диагностический лог vision_error",
        ],
        [
            "Датчик двери",
            "Не меняет состояние",
            "Обрыв, залипание",
            "FSM зависает в WAIT_ENTER/CHECK_ROOM",
            "Контроль таймаутов, сигнал ALARM, сервисное окно замены датчика",
        ],
        [
            "БД SQLite",
            "lock/недоступна",
            "Файл занят/права",
            "Нет сохранения пользователей",
            "Настройка EYEGATE_DB_PATH, WAL, тест busy_timeout",
        ],
        [
            "Vision",
            "Ложный UNKNOWN",
            "Порог/свет",
            "Доступ отклонен",
            "Подбор VISION_MATCH_THRESHOLD, лог дистанций, переобучение",
        ],
        [
            "SerialBridge",
            "Нет событий",
            "COM недоступен",
            "Нет реакции на двери",
            "Проверка SENSOR_MODE=serial, лог линий, watchdog подключения",
        ],
    ]


def power_rows() -> list[list[str]]:
    return [
        ["Luckfox Pico Ultra", "5V", "5", "25"],
        ["Камера USB/CSI", "5V", "3", "15"],
        ["Замки 2 шт", "12V", "6", "72"],
        ["Датчики 2 шт", "5V", "0.1", "0.5"],
        ["ИТОГО", "-", "-", "≈112"],
    ]


def trace_rows(facts: dict) -> list[list[str]]:
    reqs = [
        ("R-1 MJPEG монитор", "/monitor", "/api/video/mjpeg", "A"),
        (
            "R-2 Имя/UNKNOWN",
            "/monitor",
            "GateController.snapshot recognized_user_ids",
            "A/E",
        ),
        ("R-3 Автодоводчики", "config DOOR*_AUTO_CLOSE_SEC", "hw/simulated.py", "D"),
        ("R-4 SerialBridge", "hw/serial_bridge.py", "/api/sim/sensor", "F"),
        ("R-5 Персистентность БД", "db/base.py EYEGATE_DB_PATH", "db/init_db.py", "B"),
        ("R-6 FSM", "gate/fsm.py", "/ws/status", "—"),
    ]
    rows = []
    for r in reqs:
        rows.append([r[0], r[1], r[2], r[3]])
    return rows


def build_rpz() -> None:
    facts = load_facts()
    doc = copy_template()
    insert_toc(doc)

    doc.add_paragraph("Расчетно-пояснительная записка (РПЗ)", style="Title")
    if not TEMPLATE.exists():
        doc.add_paragraph(
            f"Примечание: шаблон {TEMPLATE} отсутствует, использован базовый стиль python-docx."
        )
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS_PATH.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Выбор прототипа", level=1)
    doc.add_paragraph(
        "Прототип: двухдверный mantrap с камерой и распознаванием лица/силуэта. Выбор обусловлен требованиями контроля одиночного доступа."
    )
    doc.add_paragraph(
        "Сравнение: однодверные СКУД не обеспечивают проверку второго рубежа; манжетные турникеты не поддерживают камеру в камере."
    )

    doc.add_heading("2 Анализ технического задания", level=1)
    doc.add_paragraph(
        "Ключевые требования: MJPEG /api/video/mjpeg, отображение имени/UNKNOWN, FSM IDLE→WAIT_ENTER→CHECK_ROOM→ACCESS_GRANTED/ALARM/RESET, WS /ws/status, персистентность БД SQLite, авто-закрытие дверей, SerialBridge."
    )
    doc.add_paragraph(
        "Риски: незавершенная интеграция Proteus/SerialBridge (F), возможные потери БД при lock (B), отсутствие камеры в симуляции двери 2 (C)."
    )

    doc.add_heading("3 Элементная база", level=1)
    doc.add_paragraph(
        "В текущей версии отсутствует детализированная номенклатура аппаратных компонентов (Luckfox Pico Ultra, модель камеры, датчики)."
    )
    doc.add_paragraph(
        "План внедрения: уточнить модели камеры (CSI/USB), датчиков (геркон/оптопара), замков (нормально-замкнутый), питание 12V/5V. Добавить datasheet в 99_Приложения_и_материалы."
    )

    doc.add_heading("4 Обоснование технических решений", level=1)
    doc.add_paragraph(
        "Backend: FastAPI + SQLite (WAL), модульная структура server/api/*, gate/*, db/*, vision/*, policy/*. Vision dummy/real OpenCV. WS /ws/status."
    )
    doc.add_paragraph(
        "Frontend: React/Vite SPA, маршруты /kiosk /monitor /sim /admin /enroll. Monitor использует MJPEG компонент MjpegStream, Sim управляет дверями."
    )
    doc.add_paragraph(
        "FSM: gate/fsm.py чистый автомат, gate/controller.py обрабатывает действия, policy.access решает разрешение."
    )
    doc.add_paragraph(
        "БД: таблицы users/events/settings; init_db.py создает/мигрирует и seed admin/demo."
    )

    doc.add_heading("5 Расчет надежности (оценочно)", level=1)
    doc.add_paragraph("FMEA-таблица (оценка отказов без количественных MTBF):")
    add_table(doc, ["Узел", "Отказ", "Причина", "Последствие", "Меры"], fmea_rows())
    doc.add_paragraph(
        "Допущение: средняя наработка камеры/датчиков зависит от конкретной модели; требуется уточнение после выбора элементной базы."
    )

    doc.add_heading("6 Энергопотребление и охлаждение (оценочно)", level=1)
    add_table(doc, ["Узел", "U, V", "I, A", "P, Вт"], power_rows())
    doc.add_paragraph(
        "Итого ~112 Вт (грубая оценка). Требуется уточнить после выбора конкретных компонентов; предполагается пассивное охлаждение при естественной конвекции."
    )

    doc.add_heading("7 Моделирование", level=1)
    doc.add_paragraph(
        "В текущей версии отсутствует полноценное описание модели; шаг 7 подготовит документ. Симуляция дверей реализована в hw/simulated.py и /api/sim, но камера/FOV возле двери 2 не отображается (проблема C)."
    )
    doc.add_paragraph(
        "План: добавить превью MJPEG и схему поля зрения в SimPage; подготовить сценарии аварий/предельных режимов."
    )

    doc.add_heading("8 Выводы", level=1)
    doc.add_paragraph(
        "Архитектура подтверждает выполнимость требований: REST/WS, FSM, MJPEG, SQLite. Основные доработки: устранить legacy getUserMedia, закрепить стабильный EYEGATE_DB_PATH, добавить симуляцию камеры/датчиков и диагностику vision."
    )

    doc.add_heading("9 Приложения", level=1)
    doc.add_heading("Приложение A — API (из repo_facts.json)", level=2)
    api_rows = []
    for e in facts.get("api_endpoints", []):
        api_rows.append(
            [
                e.get("method", ""),
                e.get("path", ""),
                e.get("summary", ""),
                e.get("source", ""),
            ]
        )
    if api_rows:
        add_table(doc, ["Method", "Path", "Summary", "Source"], api_rows)
    else:
        doc.add_paragraph("API: В текущей версии отсутствует (нужно заполнить).")

    doc.add_heading("Приложение B — FSM", level=2)
    doc.add_paragraph(
        "Состояния: IDLE, WAIT_ENTER, CHECK_ROOM, ACCESS_GRANTED, ACCESS_DENIED, ALARM, RESET. См. gate/fsm.py."
    )

    doc.add_heading("Приложение C — Трассируемость требований", level=2)
    add_table(
        doc,
        ["Требование", "Компонент", "Источник данных", "Проблема A–F"],
        trace_rows(facts),
    )

    doc.add_heading("Приложение D — ENV", level=2)
    env_rows = []
    for k, v in (facts.get("env_vars") or {}).items():
        env_rows.append([k, str(v)])
    if env_rows:
        add_table(doc, ["ENV", "Значение по умолчанию"], env_rows)
    else:
        doc.add_paragraph("ENV: В текущей версии отсутствует — см. .env.example.")

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    TARGET.unlink(missing_ok=True)
    doc.save(TARGET)
    print(f"Saved {TARGET}")


def main() -> None:
    build_rpz()


if __name__ == "__main__":
    main()
