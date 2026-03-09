from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
MODEL_DIR = PROJECT_ROOT / "КП_2025_EyeGate_Mantrap" / "05_Моделирование"
REPO_FACTS = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "repo_facts.json"
)
TEMPLATE = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")


def load_facts() -> dict:
    if REPO_FACTS.exists():
        try:
            return json.loads(REPO_FACTS.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def insert_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def copy_template(dest: Path) -> Document:
    if TEMPLATE.exists():
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(TEMPLATE, dest)
        return Document(dest)
    return Document()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def build_model_doc(facts: dict) -> None:
    dest = MODEL_DIR / "Описание_модели.docx"
    doc = copy_template(dest)
    insert_toc(doc)
    doc.add_paragraph("Описание модели (плейсхолдер)", style="Title")
    if not TEMPLATE.exists():
        doc.add_paragraph(f"Шаблон {TEMPLATE} отсутствует, используется базовый стиль.")
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Тип модели", level=1)
    doc.add_paragraph(
        "Комбинированная: программный симулятор дверей/замков/датчиков (hw/simulated.py, /api/sim) + план интеграции Proteus/SerialBridge (COMPIM)."
    )

    doc.add_heading("2 Состав стенда", level=1)
    doc.add_paragraph(
        "Дверь 1, дверь 2, датчики open/closed, автодоводчики (DOOR1/2_AUTO_CLOSE_SEC), камера (MJPEG), контроллер FastAPI, БД SQLite, WS /ws/status."
    )

    doc.add_heading("3 Среда моделирования", level=1)
    doc.add_paragraph(
        "SENSOR_MODE=sim — виртуальные двери. SENSOR_MODE=serial — прием строк D1/D2 или JSON. Proteus/COMPIM: план — подключение SerialBridge к виртуальному COM, отправка событий."
    )

    doc.add_heading("4 Режимы", level=1)
    add_table(
        doc,
        ["Режим", "Описание", "Критерий"],
        [
            ["Штатный", "1 известный → Door2", "Match OK, people_count=1, нет ALARM"],
            [
                "Аварийный",
                "Camera down / vision_error",
                "ALARM или deny, сообщение CAMERA DOWN",
            ],
            ["Предельный", "2 человека при закрытых дверях", "ALARM on, deny"],
            ["Unknown", "UNKNOWN лицо/силуэт", "ACCESS_DENIED или ALARM"],
        ],
    )

    doc.add_heading("5 Контрольные примеры", level=1)
    doc.add_paragraph(
        "1) PIN известного → Door1 open → CHECK_ROOM → Door2 open (ACCESS_GRANTED)."
    )
    doc.add_paragraph("2) Unknown → Door1 lock → ACCESS_DENIED.")
    doc.add_paragraph("3) >1 человек при закрытых дверях → ALARM.")
    doc.add_paragraph(
        "4) SENSOR_MODE=serial: отправить 'D1:OPEN', 'D1:CLOSED' — проверить FSM переходы."
    )

    doc.add_heading("6 Ограничения модели", level=1)
    doc.add_paragraph(
        "Нет реальной схемы/датчиков/камеры; отсутствует визуализация камеры для двери 2 в /sim (план доработки)."
    )

    doc.add_heading("7 Отличие от серийного изделия", level=1)
    doc.add_paragraph(
        "Нет аппаратного контроллера Luckfox Pico Ultra; нет реальных замков/датчиков; vision dummy может использоваться вместо реальной камеры."
    )

    doc.save(dest)


def main() -> None:
    MODEL_DIR.mkdir(parents=True, exist_ok=True)
    facts = load_facts()
    build_model_doc(facts)
    print("Built modeling doc (плейсхолдер).")


if __name__ == "__main__":
    main()
