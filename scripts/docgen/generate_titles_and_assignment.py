from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
TARGET_DIR = ROOT / "КП_2025_EyeGate_Mantrap" / "00_Админка_и_Титулы"
REPO_FACTS_REL = Path("../99_Приложения_и_материалы/repo_facts.json")
TEMPLATE_PATH = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")


def insert_toc(doc: Document) -> None:
    """Insert TOC field; update in Word after opening."""
    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def add_table(
    doc: Document, headers: list[str], rows: list[tuple[str, str, str]]
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val


def build_titles() -> None:
    doc = Document()
    insert_toc(doc)
    doc.add_paragraph(
        "Титульные листы пакета документов КП «EyeGate Mantrap»", style="Title"
    )
    if not TEMPLATE_PATH.exists():
        doc.add_paragraph(
            f"Шаблон {TEMPLATE_PATH} не найден. Использован базовый стиль python-docx."
        )
    else:
        doc.add_paragraph(
            f"Шаблон {TEMPLATE_PATH} найден. Перенесите стили вручную при необходимости."
        )
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS_REL}")
    doc.add_paragraph("Заполнители: {ФИО}, {ГРУППА}, {РУКОВОДИТЕЛЬ}, {ГОД}, {ГОРОД}.")

    doc.add_paragraph("Перечень визируемых документов", style="Heading 1")
    rows = [
        (
            "Техническое задание",
            "01_ТЗ/ТЗ_EyeGate_Mantrap.docx",
            "{Заказчик / Исполнитель}",
        ),
        (
            "РПЗ (пояснительная записка)",
            "02_КД_ЕСКД/РПЗ_EyeGate_Mantrap.docx",
            "{Руководитель / Рецензент}",
        ),
        ("ПМИ/ПСИ", "06_ПСИ_ПМИ/ПМИ_ПСИ_EyeGate_Mantrap.docx", "{Комиссия}"),
        (
            "Руководство пользователя",
            "04_ЭД_Эксплуатация/Руководство_пользователя.docx",
            "{Подпись}",
        ),
        (
            "Руководство администратора",
            "04_ЭД_Эксплуатация/Руководство_администратора.docx",
            "{Подпись}",
        ),
        (
            "Руководство программиста",
            "04_ЭД_Эксплуатация/Руководство_программиста.docx",
            "{Подпись}",
        ),
        ("Блок-схемы алгоритмов", "03_ПД_ЕСПД/Блок_схемы_алгоритмов.docx", "{Подпись}"),
    ]
    add_table(doc, ["Документ", "Файл", "Визы/подписи"], rows)
    doc.save(TARGET_DIR / "Титульники_пакет.docx")


def build_assignment() -> None:
    doc = Document()
    insert_toc(doc)
    doc.add_paragraph("Учебное задание на курсовой проект", style="Title")
    doc.add_paragraph(
        "КП «EyeGate Mantrap»: двухдверный шлюз/СКУД, FastAPI + React/Vite, MJPEG /api/video/mjpeg, WS /ws/status, FSM IDLE→WAIT_ENTER→CHECK_ROOM→ACCESS_GRANTED/ALARM/RESET."
    )
    if not TEMPLATE_PATH.exists():
        doc.add_paragraph(
            f"Шаблон {TEMPLATE_PATH} не найден. Использован базовый стиль python-docx."
        )
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS_REL}")

    add_table(
        doc,
        ["Поле", "Значение"],
        [
            ("Тема", "EyeGate Mantrap — распознавание лица+силуэта, двухдверный шлюз"),
            ("Студент", "{ФИО}"),
            ("Группа", "{ГРУППА}"),
            ("Руководитель", "{РУКОВОДИТЕЛЬ}"),
            (
                "Цель",
                "Подготовка комплекта КП 2025, MJPEG мониторинг, сохранение пользователей в SQLite, симуляция датчиков",
            ),
            (
                "Исходные данные",
                "repo_facts.json; API FastAPI; SPA /kiosk /monitor /sim /admin /enroll",
            ),
            ("Срок выполнения", "{СРОКИ}"),
            (
                "Контрольные точки",
                "Титулы → ТЗ → РПЗ → алгоритмы → руководства → ПМИ → артефакты сборки",
            ),
        ],
    )

    doc.add_paragraph("План внедрения недостающих данных", style="Heading 1")
    doc.add_paragraph(
        f"1) Получить фирменный шаблон ({TEMPLATE_PATH}) и перенести стили в титулы."
    )
    doc.add_paragraph(
        "2) Уточнить ФИО/группу/руководителя и обновить все титулы/задание."
    )
    doc.add_paragraph(
        "3) При изменении API/FSM/DB обновить repo_facts.json и титульные реквизиты."
    )

    doc.save(TARGET_DIR / "Учебное_задание_КП.docx")


def main() -> None:
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    build_titles()
    build_assignment()
    print("Generated титулы и учебное задание.")


if __name__ == "__main__":
    main()
