from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
OPS_DIR = PROJECT_ROOT / "КП_2025_EyeGate_Mantrap" / "04_ЭД_Эксплуатация"
REPO_FACTS = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "repo_facts.json"
)
TEMPLATE = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")


def load_facts() -> dict:
    if REPO_FACTS.exists():
        try:
            return json.loads(REPO_FACTS.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def insert_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def copy_template(dest: Path) -> Document:
    if TEMPLATE.exists():
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(TEMPLATE, dest)
        return Document(dest)
    return Document()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def build_user_guide(facts: dict) -> None:
    dest = OPS_DIR / "Руководство_пользователя.docx"
    doc = copy_template(dest)
    insert_toc(doc)
    doc.add_paragraph("Руководство пользователя (плейсхолдер)", style="Title")
    if not TEMPLATE.exists():
        doc.add_paragraph(f"Шаблон {TEMPLATE} отсутствует, используется базовый стиль.")
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Назначение", level=1)
    doc.add_paragraph(
        "Обеспечить проход через шлюз по PIN/карте/распознаванию, следуя инструкциям на страницах /kiosk и /monitor."
    )

    doc.add_heading("2 Интерфейс", level=1)
    doc.add_paragraph(
        "/kiosk — ввод PIN (4 цифры), статус дверей/vision. /monitor — просмотр MJPEG /api/video/mjpeg, overlay имен/UNKNOWN."
    )
    doc.add_paragraph(
        "Не предоставляйте доступ браузеру к камере — монитор использует серверный поток MJPEG."
    )

    doc.add_heading("3 Действия", level=1)
    doc.add_paragraph(
        "1) Ввести PIN на /kiosk. 2) Дождаться открытия Door1, войти. 3) Дождаться проверки комнаты (CHECK_ROOM). 4) При разрешении — открыть Door2."
    )

    doc.add_heading("4 Аварийные ситуации", level=1)
    doc.add_paragraph(
        "CAMERA DOWN — обратиться к администратору; ALARM — выйти и вызвать оператора; UNKNOWN — отказ в доступе."
    )

    doc.save(dest)


def build_admin_guide(facts: dict) -> None:
    dest = OPS_DIR / "Руководство_администратора.docx"
    doc = copy_template(dest)
    insert_toc(doc)
    doc.add_paragraph("Руководство администратора (плейсхолдер)", style="Title")
    if not TEMPLATE.exists():
        doc.add_paragraph(f"Шаблон {TEMPLATE} отсутствует, используется базовый стиль.")
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Доступ", level=1)
    doc.add_paragraph(
        "Админ логин/пароль из ENV: ADMIN_LOGIN/ADMIN_PASS. Авторизация через /api/auth/admin/login или UI /admin."
    )

    doc.add_heading("2 Управление пользователями", level=1)
    doc.add_paragraph(
        "/admin: создать пользователя (login/pin/name), кнопка Capture face → /api/users/{id}/enroll."
    )
    doc.add_paragraph("/enroll: самозапись лица (login/password → token → Capture).")

    doc.add_heading("3 Настройки", level=1)
    env_rows = []
    for k, v in (facts.get("env_vars") or {}).items():
        env_rows.append([k, str(v)])
    if env_rows:
        add_table(doc, ["ENV", "Значение"], env_rows)

    doc.add_heading("4 Мониторинг", level=1)
    doc.add_paragraph(
        "/monitor — MJPEG, WS статус. Проверить отображение имени/UNKNOWN, ALARM при >1 человека при закрытых дверях."
    )

    doc.add_heading("5 Проблемы A–F и действия", level=1)
    doc.add_paragraph(
        "A) Убрать legacy getUserMedia, проверять /api/video/mjpeg. B) EYEGATE_DB_PATH фиксировать на постоянный путь. C) Добавить камеру/поле зрения в /sim. D) Настроить DOOR*_AUTO_CLOSE_SEC. E) Логи/пороги Vision, метки UNKNOWN. F) SerialBridge протокол D1/D2 open/closed."
    )

    doc.save(dest)


def build_dev_guide(facts: dict) -> None:
    dest = OPS_DIR / "Руководство_программиста.docx"
    doc = copy_template(dest)
    insert_toc(doc)
    doc.add_paragraph("Руководство программиста (плейсхолдер)", style="Title")
    if not TEMPLATE.exists():
        doc.add_paragraph(f"Шаблон {TEMPLATE} отсутствует, используется базовый стиль.")
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Сборка и запуск", level=1)
    doc.add_paragraph(
        "Backend: `pip install -r requirements.txt`; `uvicorn server.main:app --reload`."
    )
    doc.add_paragraph(
        "Frontend: `cd web/app && npm install && npm run dev` (или build + uvicorn static)."
    )

    doc.add_heading("2 Точки интеграции", level=1)
    doc.add_paragraph(
        "MJPEG: server/api/video.py; WS: server/ws.py + status_bus; Sensors: hw/serial_bridge.py + config SENSOR_MODE/PORT/BAUD; FSM: gate/fsm.py; Controller: gate/controller.py."
    )

    doc.add_heading("3 Проблемы A–F", level=1)
    doc.add_paragraph(
        "A) Legacy /monitor getUserMedia — удалить, использовать MjpegStream. B) EYEGATE_DB_PATH стабилизировать, убрать fallback. C) Добавить камеру/FOV в симуляцию (SimPage). D) UI/реал для DOOR*_AUTO_CLOSE_SEC. E) Vision пороги/логи/UNKNOWN. F) SerialBridge инструкция и тест."
    )

    doc.add_heading("4 Тесты/планы", level=1)
    doc.add_paragraph(
        "Добавить интеграционные тесты: MJPEG доступен без getUserMedia; DB persistence при рестарте; WS /ws/status отправляет снапшот; SerialBridge парсинг D1/D2; авто-close дверей."
    )

    doc.save(dest)


def main() -> None:
    OPS_DIR.mkdir(parents=True, exist_ok=True)
    facts = load_facts()
    build_user_guide(facts)
    build_admin_guide(facts)
    build_dev_guide(facts)
    print("Built ops docs: пользователь, администратор, программист.")


if __name__ == "__main__":
    main()
