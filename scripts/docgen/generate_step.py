from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path
from typing import Iterable

PROJECT_ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")
REPO_FACTS_PATH = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "repo_facts.json"
)


def load_repo_facts() -> dict:
    if REPO_FACTS_PATH.exists():
        try:
            return json.loads(REPO_FACTS_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def ensure_template() -> bool:
    try:
        import docx  # noqa: F401
    except ImportError as exc:  # pragma: no cover - dependency hint
        raise SystemExit(
            "python-docx is required. Install via: pip install python-docx"
        ) from exc
    return TEMPLATE.exists()


def copy_template(dest: Path):
    from docx import Document

    dest.parent.mkdir(parents=True, exist_ok=True)
    if TEMPLATE.exists():
        shutil.copy(TEMPLATE, dest)
        doc = Document(dest)
    else:
        doc = Document()
        doc.add_paragraph(
            f"Шаблон {TEMPLATE} не найден. Использован базовый стиль docx."
        )
    _clear_body(doc)
    return doc


def _clear_body(doc) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def insert_toc(doc) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    r._r.append(separate)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def add_placeholder_content(doc, title: str, notes: Iterable[str]) -> None:
    doc.add_paragraph(title, style="Heading 1")
    for line in notes:
        doc.add_paragraph(line)


def build_step(step: int) -> int:
    has_template = ensure_template()
    facts = load_repo_facts()
    step_docs = {
        1: [
            (
                "КП_2025_EyeGate_Mantrap/00_Админка_и_Титулы/Титульники_пакет.docx",
                "Титульные листы (пакет)",
            ),
            (
                "КП_2025_EyeGate_Mantrap/00_Админка_и_Титулы/Учебное_задание_КП.docx",
                "Учебное задание на КП",
            ),
        ],
        2: [
            (
                "КП_2025_EyeGate_Mantrap/01_ТЗ/ТЗ_EyeGate_Mantrap.docx",
                "Техническое задание",
            ),
        ],
        3: [
            (
                "КП_2025_EyeGate_Mantrap/02_КД_ЕСКД/РПЗ_EyeGate_Mantrap.docx",
                "Расчётно-пояснительная записка",
            ),
        ],
        5: [
            (
                "КП_2025_EyeGate_Mantrap/03_ПД_ЕСПД/Блок_схемы_алгоритмов.docx",
                "Блок-схемы алгоритмов",
            ),
            (
                "КП_2025_EyeGate_Mantrap/03_ПД_ЕСПД/Описание_программы.docx",
                "Описание программы",
            ),
            (
                "КП_2025_EyeGate_Mantrap/03_ПД_ЕСПД/Исходные_тексты_программы.docx",
                "Исходные тексты программы",
            ),
        ],
        6: [
            (
                "КП_2025_EyeGate_Mantrap/04_ЭД_Эксплуатация/Руководство_пользователя.docx",
                "Руководство пользователя",
            ),
            (
                "КП_2025_EyeGate_Mantrap/04_ЭД_Эксплуатация/Руководство_администратора.docx",
                "Руководство администратора",
            ),
            (
                "КП_2025_EyeGate_Mantrap/04_ЭД_Эксплуатация/Руководство_программиста.docx",
                "Руководство программиста",
            ),
        ],
        7: [
            (
                "КП_2025_EyeGate_Mantrap/05_Моделирование/Описание_модели.docx",
                "Описание модели",
            ),
        ],
        8: [
            (
                "КП_2025_EyeGate_Mantrap/06_ПСИ_ПМИ/ПМИ_ПСИ_EyeGate_Mantrap.docx",
                "Программа и методика испытаний",
            ),
        ],
    }

    targets = step_docs.get(step, [])
    if not targets:
        print(f"No document generation configured for step {step}.")
        return 1

    for path_str, title in targets:
        dest = PROJECT_ROOT / path_str
        doc = copy_template(dest)
        insert_toc(doc)
        notes = [
            "Плейсхолдер. Наполнить содержимым по КП_2025.",
            f"Источник фактов: {REPO_FACTS_PATH.relative_to(PROJECT_ROOT)}",
        ]
        if not has_template:
            notes.append(f"Внимание: шаблон {TEMPLATE} не найден, базовый стиль docx.")
        if facts:
            notes.append(f"Доступны факты: {', '.join(sorted(facts.keys()))}")
        add_placeholder_content(doc, title, notes)
        doc.save(dest)
        print(f"Generated placeholder: {dest}")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate KP_2025 documents from template with step-aware scaffolding."
    )
    parser.add_argument("--step", type=int, required=True, help="Step number (0-9)")
    args = parser.parse_args()

    if args.step == 0:
        print("Step 0: only preparation; documents are not generated here.")
        return 0

    return build_step(args.step)


if __name__ == "__main__":
    sys.exit(main())
