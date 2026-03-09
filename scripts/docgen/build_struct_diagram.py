from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DIAGRAM_PNG = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "02_КД_ЕСКД"
    / "Схема_структурная_системы.png"
)
DIAGRAM_SRC = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "diagrams_source"
    / "struct_scheme.mmd"
)
SPEC_DOC = PROJECT_ROOT / "КП_2025_EyeGate_Mantrap" / "02_КД_ЕСКД" / "Спецификация.docx"
PE_DOC = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "02_КД_ЕСКД"
    / "ПЭ_перечень_элементов.docx"
)
SCHEMA_PLACEHOLDER = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "02_КД_ЕСКД"
    / "Схема_электрическая_принципиальная_PLACEHOLDER.pdf"
)
REPO_FACTS = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "repo_facts.json"
)
TEMPLATE = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")


def insert_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def ensure_struct_diagram() -> None:
    DIAGRAM_PNG.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)

    def box(x, y, w, h, text):
        rect = plt.Rectangle(
            (x, y), w, h, linewidth=1.5, edgecolor="#2563eb", facecolor="#e0e7ff"
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9)

    box(0.4, 2.2, 2.0, 1.6, "Дверь 1\nдатчик/замок")
    box(7.6, 2.2, 2.0, 1.6, "Дверь 2\nдатчик/замок")
    box(3.4, 2.0, 3.0, 2.0, "Камера\nMJPEG /api/video/mjpeg")
    box(3.4, 0.6, 3.0, 1.0, "Vision\n(детекция лица+силуэт)")
    box(3.4, 4.3, 3.0, 1.0, "Контроллер\nFastAPI + FSM")
    box(3.4, 5.5, 3.0, 0.7, "WS /ws/status\nREST /api/*")
    box(0.4, 0.6, 2.0, 1.0, "SerialBridge\n(SENSOR_MODE=serial)")
    box(7.6, 0.6, 2.0, 1.0, "DB SQLite\nusers/events")

    arrowprops = dict(arrowstyle="->", color="#111827", linewidth=1.2)
    ax.annotate("", xy=(3.4, 2.8), xytext=(2.4, 3.0), arrowprops=arrowprops)
    ax.annotate("", xy=(6.4, 2.8), xytext=(7.6, 3.0), arrowprops=arrowprops)
    ax.annotate(
        "REST/WS", xy=(5.0, 5.5), xytext=(5.0, 4.3), ha="center", arrowprops=arrowprops
    )
    ax.annotate(
        "Frames", xy=(5.0, 3.5), xytext=(5.0, 4.0), ha="center", arrowprops=arrowprops
    )
    ax.annotate(
        "Analysis", xy=(5.0, 2.0), xytext=(5.0, 1.6), ha="center", arrowprops=arrowprops
    )
    ax.annotate(
        "Sensors", xy=(1.4, 1.6), xytext=(1.4, 2.2), ha="center", arrowprops=arrowprops
    )
    ax.annotate(
        "Events", xy=(8.6, 1.6), xytext=(8.6, 2.2), ha="center", arrowprops=arrowprops
    )

    fig.savefig(DIAGRAM_PNG, dpi=180, bbox_inches="tight")
    plt.close(fig)

    DIAGRAM_SRC.parent.mkdir(parents=True, exist_ok=True)
    DIAGRAM_SRC.write_text(
        """flowchart LR
  Door1[Дверь 1: датчик/замок] -- события --> Controller
  Door2[Дверь 2: датчик/замок] -- события --> Controller
  Camera -- MJPEG /api/video/mjpeg --> SPA_Monitor
  Controller -- WS /ws/status --> SPA_Monitor
  Controller -- REST /api/* --> SPA_Admin
  Controller -- SQLite --> DB[(DB users/events)]
  SerialBridge -- D1/D2 JSON --> Controller
""",
        encoding="utf-8",
    )


def build_spec_doc() -> None:
    doc = Document()
    insert_toc(doc)
    doc.add_paragraph("Спецификация (плейсхолдер)", style="Title")
    doc.add_paragraph(
        "В текущей версии отсутствует детализированная спецификация по ЕСКД."
    )
    doc.add_paragraph("План: заполнить после выбора компонентной базы.")
    add_table(
        doc,
        ["Поз.", "Обозначение", "Наименование", "Кол.", "Примечание"],
        [
            [
                "A1",
                "CTRL",
                "Контроллер (Luckfox Pico Ultra, план)",
                "1",
                "Требуется уточнить модель/питание",
            ],
            ["A2", "CAM1", "Камера (USB/CSI)", "1", "Требуется модель/интерфейс"],
            [
                "A3",
                "LOCK1",
                "Замок электромагнитный Door1",
                "1",
                "НЗ/НР, питание 12V — уточнить",
            ],
            [
                "A4",
                "LOCK2",
                "Замок электромагнитный Door2",
                "1",
                "НЗ/НР, питание 12V — уточнить",
            ],
            [
                "A5",
                "SENS1",
                "Датчик двери 1 (геркон/оптопара)",
                "1",
                "Тип/контакт — уточнить",
            ],
            [
                "A6",
                "SENS2",
                "Датчик двери 2 (геркон/оптопара)",
                "1",
                "Тип/контакт — уточнить",
            ],
            ["A7", "PSU", "Блок питания 12V/5V", "1", "Мощность/разъемы — уточнить"],
            [
                "A8",
                "IF1",
                "Интерфейс SerialBridge (UART/USB)",
                "1",
                "COM параметры — уточнить",
            ],
            ["A9", "CAB", "Кабели/жгуты", "н/д", "Длины/типы — уточнить"],
        ],
    )
    doc.add_paragraph("Документ остается плейсхолдером до внесения реальных данных.")
    doc.save(SPEC_DOC)


def build_pe_doc() -> None:
    doc = Document()
    insert_toc(doc)
    doc.add_paragraph("Перечень элементов (ПЭ, плейсхолдер)", style="Title")
    doc.add_paragraph("В текущей версии отсутствуют реальные схемы/номенклатура.")
    add_table(
        doc,
        ["Поз.", "Наименование", "Кол.", "Примечание"],
        [
            ["A1", "Контроллер (Luckfox Pico Ultra)", "1", "Планируемый вычислитель"],
            ["A2", "Камера (USB/CSI)", "1", "Требуется модель"],
            ["A3", "Замок Door1", "1", "Тип/ток потребления — уточнить"],
            ["A4", "Замок Door2", "1", "Тип/ток потребления — уточнить"],
            ["A5", "Датчик Door1", "1", "NO/NC — уточнить"],
            ["A6", "Датчик Door2", "1", "NO/NC — уточнить"],
            ["A7", "Блок питания 12V", "1", "Мощность — уточнить"],
            ["A8", "Блок питания 5V", "1", "Мощность — уточнить"],
            ["A9", "Кабели/разъемы", "н/д", "Подобрать по трассировке"],
        ],
    )
    doc.add_paragraph(
        "План: после подготовки KiCad .kicad_sch/.kicad_pcb экспортировать актуальный ПЭ."
    )
    doc.save(PE_DOC)


def build_schema_placeholder() -> None:
    SCHEMA_PLACEHOLDER.parent.mkdir(parents=True, exist_ok=True)
    if not SCHEMA_PLACEHOLDER.exists():
        SCHEMA_PLACEHOLDER.write_bytes(
            b"PLACEHOLDER: electrical schematic to be added (KiCad/PDF)."
        )


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def main() -> None:
    ensure_struct_diagram()
    build_spec_doc()
    build_pe_doc()
    build_schema_placeholder()
    print(f"Saved structural diagram {DIAGRAM_PNG}")
    print(f"Saved spec {SPEC_DOC}")
    print(f"Saved ПЭ {PE_DOC}")
    print(f"Placeholder schematic {SCHEMA_PLACEHOLDER}")


if __name__ == "__main__":
    main()
