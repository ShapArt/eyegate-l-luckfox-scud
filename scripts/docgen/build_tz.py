from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
TARGET = PROJECT_ROOT / "КП_2025_EyeGate_Mantrap" / "01_ТЗ" / "ТЗ_EyeGate_Mantrap.docx"
REPO_FACTS_PATH = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "repo_facts.json"
)
TEMPLATE = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")


def load_facts() -> dict:
    if REPO_FACTS_PATH.exists():
        try:
            return json.loads(REPO_FACTS_PATH.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def copy_template() -> Document:
    if TEMPLATE.exists():
        TARGET.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(TEMPLATE, TARGET)
        doc = Document(TARGET)
        clear_body(doc)
    else:
        doc = Document()
    return doc


def insert_toc(doc: Document) -> None:
    """Add TOC field; update in Word."""
    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def build_requirement_rows() -> list[list[str]]:
    return [
        [
            "R-1",
            "Монитор отображает MJPEG из /api/video/mjpeg (без getUserMedia)",
            "A",
            "Открыть /monitor → запрос /api/video/mjpeg; камера браузера не запрашивается",
            "frontend MonitorPage.tsx; server/api/video.py; tests TBD",
        ],
        [
            "R-2",
            "Отображать имя/логин владельца, иначе красный UNKNOWN; силует+лицо",
            "A",
            "WS /ws/status содержит recognized_user_ids; Monitor показывает имя/UNKNOWN",
            "MonitorPage.tsx; GateController.snapshot; tests TBD",
        ],
        [
            "R-3",
            "Автодоводчики дверей 1 и 2 (таймеры)",
            "B",
            "Установить DOOR1/DOOR2_AUTO_CLOSE_SEC >0 → двери закрываются автоматически",
            "server/config.py; hw/simulated.py; tests TBD",
        ],
        [
            "R-4",
            "Сенсоры/Proteus: SerialBridge протокол (D1:OPEN/CLOSED; JSON) готов к интеграции",
            "B",
            "SENSOR_MODE=serial; принимаются строки D1/D2 или JSON",
            "hw/serial_bridge.py; server/deps.start_gate_controller; tests TBD",
        ],
        [
            "R-5",
            "Персистентность пользователей в SQLite сохраняется после перезапуска",
            "A",
            "EYEGATE_DB_PATH указывает на data/eyegate_scud.db; пользователи сохраняются",
            "db/base.py; db/init_db.py; tests TBD",
        ],
        [
            "R-6",
            "FSM шлюза IDLE→WAIT_ENTER→CHECK_ROOM→ACCESS_GRANTED/ALARM/RESET",
            "A",
            "GateController.snapshot.state отражает переходы по событиям",
            "gate/fsm.py; gate/controller.py; tests TBD",
        ],
        [
            "R-7",
            "WS статус /ws/status и REST /api/status доступны",
            "A",
            "Подключение WS: первое сообщение со снапшотом; GET /api/status возвращает JSON",
            "server/ws.py; server/api/status.py",
        ],
        [
            "R-8",
            "SPA страницы: /kiosk, /monitor, /sim, /admin, /enroll",
            "B",
            "Маршруты Vite/React активны; deep-link открывается",
            "web/app/src/App.tsx",
        ],
        [
            "R-9",
            "Видеокадр у двери 2 в симуляции (камера/поле зрения) — план внедрения",
            "C",
            "В текущей версии отсутствует; подготовить схему и UI обновление",
            "web/app/src/pages/SimPage.tsx (план)",
        ],
        [
            "R-10",
            "Диагностика распознавания лиц: пороги/логи/UNKNOWN",
            "B",
            "VISION_* из .env управляют порогами; UI показывает match/distance/UNKNOWN",
            "server/config.py; MonitorPage.tsx",
        ],
    ]


def build_api_rows(facts: dict) -> list[list[str]]:
    items = facts.get("api_endpoints") or []
    rows: list[list[str]] = []
    for e in items:
        rows.append(
            [
                e.get("method", ""),
                e.get("path", ""),
                e.get("summary", ""),
                e.get("source", ""),
            ]
        )
    return rows


def build_env_rows(facts: dict) -> list[list[str]]:
    env = facts.get("env_vars") or {}
    rows = []
    for k, v in env.items():
        rows.append([k, str(v)])
    return rows


def build_tz() -> None:
    facts = load_facts()
    doc = copy_template()
    insert_toc(doc)

    doc.add_paragraph("Техническое задание", style="Title")
    if not TEMPLATE.exists():
        doc.add_paragraph(
            f"Примечание: шаблон {TEMPLATE} отсутствует, использован базовый стиль python-docx."
        )
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS_PATH.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Общие сведения", level=1)
    doc.add_paragraph(
        "Проект: EyeGate Mantrap — двухдверный шлюз/СКУД на FastAPI (backend) и React/Vite (SPA)."
    )
    doc.add_paragraph(
        "Функции: MJPEG видеопоток /api/video/mjpeg, WebSocket /ws/status, FSM IDLE→WAIT_ENTER→CHECK_ROOM→ACCESS_GRANTED/ALARM/RESET."
    )
    doc.add_paragraph(
        "Архитектура: FastAPI + SQLite (WAL) + Vision (dummy/real OpenCV), SPA маршруты /kiosk /monitor /sim /admin /enroll."
    )

    doc.add_heading("2 Назначение и цели", level=1)
    doc.add_paragraph(
        "Цель: обеспечить контроль доступа через двухдверный шлюз с проверкой лица/силуэта и журналированием событий."
    )
    doc.add_paragraph(
        "Результат: комплект КП-2025 с ТЗ, РПЗ, схемами, алгоритмами, руководствами, ПМИ, артефактами сборки."
    )

    doc.add_heading("3 Характеристика объекта автоматизации", level=1)
    doc.add_paragraph(
        "Объект: шлюз с двумя дверями, замками, датчиками открыто/закрыто, камерой наблюдения, сигнализацией."
    )
    doc.add_paragraph(
        "Каналы связи: REST/WS HTTP(S), сериализированные события сенсоров (SerialBridge) при SENSOR_MODE=serial."
    )
    doc.add_paragraph(
        "Клиентская часть: SPA Vite/React, страницы для киоска, мониторинга, симуляции, администрирования, самозаписи лица."
    )

    doc.add_heading("4 Требования к системе", level=1)
    doc.add_heading("4.1 Функциональные требования", level=2)
    doc.add_paragraph(
        "Функциональные требования собраны из ТЗ и проблем A–F. Таблица содержит критерии приемки и привязку к модулям."
    )
    add_table(
        doc,
        ["ID", "Формулировка", "Приоритет", "Критерий приемки", "Модуль/тест"],
        build_requirement_rows(),
    )

    doc.add_heading("4.2 Надежность и сохранность данных", level=2)
    doc.add_paragraph(
        "БД: SQLite, путь EYEGATE_DB_PATH (default data/eyegate_scud.db), WAL включен. Требование: данные пользователей сохраняются при перезапуске."
    )
    doc.add_paragraph(
        "При блокировках БД fallback недопустим в проде — требуется контроль доступности (план: тест WAL/locking)."
    )

    doc.add_heading("4.3 Безопасность и доступ", level=2)
    doc.add_paragraph(
        "Админ доступ: /api/auth/admin/login по ADMIN_LOGIN/ADMIN_PASS, либо X-Admin-Token (если задан)."
    )
    doc.add_paragraph(
        "Пользовательский доступ: /api/auth/login, /api/auth/pin; роли user/admin; статусы pending/active/rejected."
    )
    doc.add_paragraph(
        "Видео: запрещено getUserMedia в /monitor; только серверный MJPEG."
    )

    doc.add_heading("4.4 Интерфейсы и протоколы", level=2)
    doc.add_paragraph(
        'REST API и WS указаны в Приложении A. Сенсоры: строки D1:OPEN/CLOSED или JSON {"door":1,"state":"open"}.'
    )
    doc.add_paragraph(
        "FSM: описана в gate/fsm.py (Приложение B). WS payload: структура GateStatus (Приложение C)."
    )

    doc.add_heading("4.5 Условия эксплуатации", level=2)
    doc.add_paragraph(
        "Режимы: demo (EYEGATE_DEMO_MODE=1) и prod (EYEGATE_ENV=prod). Аппаратная цель — Luckfox Pico Ultra (плейсхолдер)."
    )

    doc.add_heading("5 Состав и содержание работ", level=1)
    doc.add_paragraph(
        "Разработка и документирование модулей backend, frontend, симуляции датчиков, интеграция Vision, подготовка ПМИ."
    )
    doc.add_paragraph(
        "Подготовка схем/алгоритмов, руководств пользователя/администратора/программиста, тест-план ПСИ."
    )

    doc.add_heading("6 Порядок контроля и приемки", level=1)
    doc.add_paragraph(
        "Контроль через ПМИ (см. STEP 8). Проверки: MJPEG поток доступен; WS работает; авто-закрытие дверей; персистентность БД; реакции на A–F кейсы."
    )

    doc.add_heading("7 Требования к документированию", level=1)
    doc.add_paragraph(
        "Документы формируются на основе repo_facts.json. Титульные листы по ГОСТ Р 6.30-97; ТЗ — ГОСТ 34.602-89/19.201-78; ПМИ — ГОСТ 19.301."
    )

    doc.add_heading("8 Источники разработки", level=1)
    doc.add_paragraph(
        "Репозиторий: текущее состояние рабочей директории. Компоненты API/SPA: см. Приложения."
    )
    doc.add_paragraph(
        "Факт-источник: repo_facts.json. Дополнительно: ГОСТы 34.602-89, 19.201-78, ГОСТ Р 51241-98."
    )

    doc.add_heading("9 Приложения", level=1)
    doc.add_heading("Приложение A — API", level=2)
    api_rows = build_api_rows(facts)
    if api_rows:
        add_table(doc, ["Method", "Path", "Summary", "Source"], api_rows)
    else:
        doc.add_paragraph(
            "API: В текущей версии отсутствует (нужно заполнить из FastAPI)."
        )

    doc.add_heading("Приложение B — FSM", level=2)
    doc.add_paragraph(
        "Состояния: IDLE, WAIT_ENTER, CHECK_ROOM, ACCESS_GRANTED, ACCESS_DENIED, ALARM, RESET."
    )
    doc.add_paragraph(
        "События: CardPresented, AuthResult, DoorClosedChanged(1/2), RoomAnalyzed, TimeoutEnter/Check/Exit/Alarm, Reset."
    )
    doc.add_paragraph(
        "Действия: LOCK/UNLOCK door1/door2, LOCK_BOTH, START_*_TIMEOUT, CANCEL_ALL_TIMEOUTS, START_ROOM_ANALYSIS, SET_ALARM_ON/OFF, LOG_EVENT, CLEAR_CONTEXT."
    )

    doc.add_heading("Приложение C — WS / GateStatus payload", level=2)
    doc.add_paragraph(
        "Поля корня: state, current_card_id, current_user_id, doors, alarm_on, last_event, vision, vision_required, demo_mode, timestamp, policy, room_samples."
    )
    doc.add_paragraph(
        "doors: door1_closed, door2_closed, lock1_unlocked, lock2_unlocked, lock1_power, lock2_power, sensor1_open, sensor2_open."
    )
    doc.add_paragraph(
        "vision: provider, people_count, boxes, faces, vision_state, last_frame_ts, fps, vision_error, match, match_distance, matched_user_id, recognized_user_ids, recognized_scores, frame_w, frame_h."
    )

    doc.add_heading("Приложение D — ENV", level=2)
    env_rows = build_env_rows(facts)
    if env_rows:
        add_table(doc, ["ENV", "Значение по умолчанию"], env_rows)
    else:
        doc.add_paragraph("ENV: В текущей версии отсутствует — см. .env.example.")

    doc.add_heading("Приложение E — Известные проблемы A–F", level=2)
    doc.add_paragraph(
        "A) Legacy /monitor использует getUserMedia — требуется MJPEG /api/video/mjpeg. План: удалить legacy static, проверить MonitorPage.tsx."
    )
    doc.add_paragraph(
        "B) Персистентность БД: fallback *_fallback.db при lock — требуется устойчивый путь EYEGATE_DB_PATH, тест WAL."
    )
    doc.add_paragraph(
        "C) Симуляция камеры у двери 2 отсутствует — требуется схема/обзор FOV и UI обновление."
    )
    doc.add_paragraph(
        "D) Авто-закрытие дверей: env DOOR1/2_AUTO_CLOSE_SEC поддержаны, но UI/док не покрыты — доработать."
    )
    doc.add_paragraph(
        "E) Распознавание лиц: нужны пороги/логи/UNKNOWN — использовать VISION_* и вывод имён/UNKNOWN."
    )
    doc.add_paragraph(
        "F) Proteus/SerialBridge: только частичная готовность — описать протокол и добавить интеграционный тест."
    )

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    TARGET.unlink(missing_ok=True)
    doc.save(TARGET)
    print(f"Saved {TARGET}")


def main() -> None:
    build_tz()


if __name__ == "__main__":
    main()
