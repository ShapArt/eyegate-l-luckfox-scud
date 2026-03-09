from __future__ import annotations

import json
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

PROJECT_ROOT = Path(__file__).resolve().parents[2]
PD_DIR = PROJECT_ROOT / "КП_2025_EyeGate_Mantrap" / "03_ПД_ЕСПД"
REPO_FACTS = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "repo_facts.json"
)
TEMPLATE = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")


def load_facts() -> dict:
    if REPO_FACTS.exists():
        try:
            return json.loads(REPO_FACTS.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def insert_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def copy_template(dest: Path) -> Document:
    if TEMPLATE.exists():
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(TEMPLATE, dest)
        return Document(dest)
    return Document()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def _build_fsm_png(path: Path) -> None:
    PD_DIR.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)

    def box(x, y, w, h, text):
        rect = plt.Rectangle(
            (x, y), w, h, linewidth=1.4, edgecolor="#0ea5e9", facecolor="#e0f2fe"
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9)

    box(1.0, 4.2, 2.0, 0.8, "IDLE")
    box(4.0, 4.2, 2.0, 0.8, "WAIT_ENTER")
    box(7.0, 4.2, 2.0, 0.8, "CHECK_ROOM")
    box(7.0, 2.6, 2.0, 0.8, "ACCESS_GRANTED")
    box(7.0, 1.0, 2.0, 0.8, "ACCESS_DENIED")
    box(4.0, 1.0, 2.0, 0.8, "ALARM")
    box(1.0, 1.0, 2.0, 0.8, "RESET")

    arrowprops = dict(arrowstyle="->", color="#0f172a", linewidth=1.2)
    ax.annotate("Card/Auth OK", xy=(3.0, 4.6), xytext=(4.0, 4.6), arrowprops=arrowprops)
    ax.annotate(
        "Door1 closed\nRoom analysis",
        xy=(6.0, 4.6),
        xytext=(7.0, 4.6),
        arrowprops=arrowprops,
    )
    ax.annotate("Policy=open", xy=(8.0, 4.2), xytext=(8.0, 3.4), arrowprops=arrowprops)
    ax.annotate("Policy=deny", xy=(8.0, 4.2), xytext=(8.0, 2.0), arrowprops=arrowprops)
    ax.annotate("Policy=alarm", xy=(6.0, 4.2), xytext=(5.0, 1.4), arrowprops=arrowprops)
    ax.annotate(
        "Timeout/reset", xy=(2.0, 4.2), xytext=(2.0, 1.8), arrowprops=arrowprops
    )
    ax.annotate("Reset", xy=(7.0, 2.6), xytext=(3.0, 1.4), arrowprops=arrowprops)

    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def build_algo_doc(facts: dict) -> None:
    dest = PD_DIR / "Блок_схемы_алгоритмов.docx"
    doc = copy_template(dest)
    insert_toc(doc)
    doc.add_paragraph("Блок-схемы алгоритмов", style="Title")
    if not TEMPLATE.exists():
        doc.add_paragraph(
            f"Шаблон {TEMPLATE} отсутствует, используется базовый стиль python-docx."
        )
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS.relative_to(PROJECT_ROOT)}")

    fsm_png = PD_DIR / "fsm_block.png"
    _build_fsm_png(fsm_png)

    doc.add_heading("1 FSM шлюза", level=1)
    doc.add_paragraph(
        "IDLE → WAIT_ENTER → CHECK_ROOM → ACCESS_GRANTED/ALARM/RESET. События: CardPresented, AuthResult, DoorClosedChanged, RoomAnalyzed, Timeout, Reset."
    )
    if fsm_png.exists():
        doc.add_picture(str(fsm_png), width=Inches(5.5))
        doc.add_paragraph("Рис. 1.1 — FSM шлюза (упрощенная блок-схема).")

    doc.add_heading("2 Сценарии A–F (проблемные ветки)", level=1)
    add_table(
        doc,
        ["ID", "Обнаружено", "Действие/проверка"],
        [
            [
                "A",
                "Legacy /monitor спрашивает камеру",
                "Убрать getUserMedia; оставить /api/video/mjpeg; в Network только MJPEG",
            ],
            [
                "B",
                "Нет персистентности БД",
                "EYEGATE_DB_PATH=data/eyegate_scud.db; тест перезапуска без потери users",
            ],
            [
                "C",
                "Нет камеры/поле зрения у двери 2 (sim)",
                "Добавить MJPEG/FOV в /sim; сейчас пометка UNKNOWN",
            ],
            [
                "D",
                "Автодоводчики не настроены",
                "DOOR1/2_AUTO_CLOSE_SEC>0; проверить закрытие таймером",
            ],
            [
                "E",
                "Ошибки распознавания лиц",
                "Логи/пороги VISION_MATCH_THRESHOLD; UNKNOWN красным, имя по user_id",
            ],
            [
                "F",
                "Proteus/SerialBridge частично",
                "Протокол D1/D2 open/closed, JSON; интеграционный тест",
            ],
        ],
    )
    doc.add_paragraph(
        "Для финала: заменить схему утвержденной, зафиксировать фактические проверки A–F."
    )
    doc.save(dest)


def build_struct_po(facts: dict) -> None:
    dest = PD_DIR / "Структурная_схема_ПО.png"
    PD_DIR.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.axis("off")
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)

    def box(x, y, w, h, text):
        rect = plt.Rectangle(
            (x, y), w, h, linewidth=1.4, edgecolor="#0891b2", facecolor="#e0f2fe"
        )
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9)

    box(0.5, 3.5, 2.4, 1.0, "API FastAPI\n/server/api/*")
    box(0.5, 2.0, 2.4, 1.0, "WS /ws/status\nstatus_bus")
    box(0.5, 0.5, 2.4, 1.0, "GateController\nFSM/policy")
    box(3.2, 3.5, 2.4, 1.0, "Vision\nOpenCV/dummy")
    box(3.2, 2.0, 2.4, 1.0, "DB SQLite\nusers/events")
    box(3.2, 0.5, 2.4, 1.0, "SerialBridge\nsim/serial sensors")
    box(6.0, 4.3, 3.0, 1.2, "SPA React/Vite\n/kiosk /monitor /sim /admin /enroll")

    arrowprops = dict(arrowstyle="->", color="#0f172a", linewidth=1.1)
    ax.annotate("", xy=(1.7, 3.5), xytext=(1.7, 3.0), arrowprops=arrowprops)
    ax.annotate("", xy=(1.7, 2.0), xytext=(1.7, 1.5), arrowprops=arrowprops)
    ax.annotate("", xy=(1.7, 0.5), xytext=(4.4, 0.5), arrowprops=arrowprops)
    ax.annotate("", xy=(4.4, 3.5), xytext=(4.4, 3.0), arrowprops=arrowprops)
    ax.annotate("", xy=(4.4, 2.0), xytext=(4.4, 1.5), arrowprops=arrowprops)
    ax.annotate("", xy=(4.4, 1.5), xytext=(4.4, 2.0), arrowprops=arrowprops)
    ax.annotate("", xy=(7.5, 4.3), xytext=(5.6, 4.0), arrowprops=arrowprops)
    ax.annotate("", xy=(5.6, 4.0), xytext=(7.5, 4.3), arrowprops=arrowprops)

    fig.savefig(dest, dpi=180, bbox_inches="tight")
    plt.close(fig)


def build_description(facts: dict) -> None:
    dest = PD_DIR / "Описание_программы.docx"
    doc = copy_template(dest)
    insert_toc(doc)
    doc.add_paragraph("Описание программы (ГОСТ 19.402-78, плейсхолдер)", style="Title")
    if not TEMPLATE.exists():
        doc.add_paragraph(f"Шаблон {TEMPLATE} отсутствует, используется базовый стиль.")
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Назначение", level=1)
    doc.add_paragraph(
        "Программа EyeGate Mantrap обеспечивает контроль доступа через двухдверный шлюз с распознаванием лица/силуэта и ведением событий."
    )

    doc.add_heading("2 Функции", level=1)
    doc.add_paragraph(
        "REST API (/api/status, /api/video/mjpeg, /api/auth, /api/users, /api/events, /api/sim), WebSocket /ws/status, FSM управления дверями, MJPEG поток, симулятор датчиков."
    )

    doc.add_heading("3 Входные данные", level=1)
    doc.add_paragraph(
        "Запросы REST/WS, события датчиков (SerialBridge), видеопоток камеры (реальной/симулятор), PIN/логин/пароль пользователей."
    )

    doc.add_heading("4 Выходные данные", level=1)
    doc.add_paragraph(
        "JSON статуса (GateStatus), MJPEG кадры, записи в БД users/events, ответы API/WS."
    )

    doc.add_heading("5 Ограничения", level=1)
    doc.add_paragraph(
        "Нет реальной принципиальной схемы/PCB; отсутствуют модели камеры/датчиков; симуляция камеры возле двери 2 не реализована."
    )

    doc.add_heading("6 Сообщения и ошибки", level=1)
    doc.add_paragraph(
        "Коды API: LOGIN_LOCKED, INVALID_CREDENTIALS, USER_PENDING, USER_BLOCKED, BAD_PIN, ENROLL_FAILED; WS ошибок нет."
    )

    doc.add_heading("7 Известные проблемы A–F", level=1)
    doc.add_paragraph(
        "A) Legacy /monitor с getUserMedia — исключить, использовать /api/video/mjpeg."
    )
    doc.add_paragraph(
        "B) Персистентность БД: риск fallback; нужен стабильный EYEGATE_DB_PATH."
    )
    doc.add_paragraph("C) Симуляция камеры/поле зрения у двери 2 отсутствует.")
    doc.add_paragraph(
        "D) Авто-доводчики: DOOR*_AUTO_CLOSE_SEC, UI/реализация требуют проверки."
    )
    doc.add_paragraph("E) Диагностика лиц: пороги/логи/UNKNOWN в UI.")
    doc.add_paragraph(
        "F) Proteus/SerialBridge: частичная готовность, требуется инструкция и тест."
    )

    doc.save(dest)


def build_sources_doc(facts: dict) -> None:
    dest = PD_DIR / "Исходные_тексты_программы.docx"
    doc = copy_template(dest)
    insert_toc(doc)
    doc.add_paragraph(
        "Исходные тексты программы (структура и ключевые листинги)", style="Title"
    )
    if not TEMPLATE.exists():
        doc.add_paragraph(f"Шаблон {TEMPLATE} отсутствует, используется базовый стиль.")
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Структура каталогов", level=1)
    tree = facts.get("repo_tree", {})
    for k, v in tree.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("2 Ключевые файлы", level=1)
    doc.add_paragraph(
        "Backend: server/main.py, server/api/*.py (status, users, events, auth, sim, video, camera_control), server/ws.py."
    )
    doc.add_paragraph(
        "FSM/контроллер: gate/fsm.py, gate/controller.py, policy/access.py."
    )
    doc.add_paragraph("DB: db/base.py, db/models.py, db/init_db.py.")
    doc.add_paragraph("Vision: vision/service.py, matcher.py, embeddings.py.")
    doc.add_paragraph(
        "Frontend: web/app/src/pages/*.tsx (Kiosk, Monitor, Sim, Admin, Enroll), components/MjpegStream.tsx, lib/api.ts."
    )

    doc.add_heading("3 Инструкция сборки", level=1)
    doc.add_paragraph(
        "Backend: `pip install -r requirements.txt`; запуск `uvicorn server.main:app --reload`."
    )
    doc.add_paragraph(
        "Frontend: `cd web/app && npm install && npm run dev` или `npm run build` + `uvicorn` для дистрибутива."
    )

    doc.add_heading("4 Примеры листингов (добавить вручную)", level=1)
    doc.add_paragraph(
        "GateController.snapshot (vision/doors), server/api/video.py (MJPEG), web/app/src/pages/MonitorPage.tsx (MJPEG UI), hw/serial_bridge.py (парсинг сенсоров)."
    )

    doc.add_heading("5 План дополнения", level=1)
    doc.add_paragraph(
        "Добавить фактические листинги к разделу 4 после фиксов проблем A–F."
    )

    doc.save(dest)


def main() -> None:
    PD_DIR.mkdir(parents=True, exist_ok=True)
    facts = load_facts()
    build_algo_doc(facts)
    build_struct_po(facts)
    build_description(facts)
    build_sources_doc(facts)
    print("Built PD documents: алгоритмы, структурная схема ПО, описание, исходники.")


if __name__ == "__main__":
    main()
