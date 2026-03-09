from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[2]
PMI_DIR = PROJECT_ROOT / "КП_2025_EyeGate_Mantrap" / "06_ПСИ_ПМИ"
REPO_FACTS = (
    PROJECT_ROOT
    / "КП_2025_EyeGate_Mantrap"
    / "99_Приложения_и_материалы"
    / "repo_facts.json"
)
TEMPLATE = Path("/mnt/data/РПЗ_СБ_Щит (1).docx")


def load_facts() -> dict:
    if REPO_FACTS.exists():
        try:
            return json.loads(REPO_FACTS.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            return {}
    return {}


def insert_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    r._r.append(instr)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r._r.append(sep)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)


def copy_template(dest: Path) -> Document:
    if TEMPLATE.exists():
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy(TEMPLATE, dest)
        return Document(dest)
    return Document()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def build_pmi(facts: dict) -> None:
    dest = PMI_DIR / "ПМИ_ПСИ_EyeGate_Mantrap.docx"
    doc = copy_template(dest)
    insert_toc(doc)
    doc.add_paragraph(
        "Программа и методика испытаний (ПМИ/ПСИ) — плейсхолдер", style="Title"
    )
    if not TEMPLATE.exists():
        doc.add_paragraph(f"Шаблон {TEMPLATE} отсутствует, используется базовый стиль.")
    doc.add_paragraph(f"Источник фактов: {REPO_FACTS.relative_to(PROJECT_ROOT)}")

    doc.add_heading("1 Цель испытаний", level=1)
    doc.add_paragraph(
        "Подтвердить выполнение требований R-1..R-10 (MJPEG, имя/UNKNOWN, FSM, автодоводчики, SerialBridge, персистентность БД, UI маршруты)."
    )

    doc.add_heading("2 Перечень испытаний", level=1)
    add_table(
        doc,
        ["ID", "Описание", "Ожидаемый результат"],
        [
            [
                "T1",
                "MJPEG /api/video/mjpeg доступен из /monitor без getUserMedia",
                "Кадры отображаются, браузер не запрашивает камеру",
            ],
            [
                "T2",
                "WS /ws/status передает снапшот после подключения",
                "Первое сообщение содержит GateStatus, далее обновления",
            ],
            [
                "T3",
                "PIN известного пользователя → Door2 открывается",
                "FSM ACCESS_GRANTED, recognized_user_ids содержит user_id",
            ],
            ["T4", "Unknown лицо → ACCESS_DENIED", "Door2 закрыта, ALARM нет"],
            ["T5", "2 человека при закрытых дверях", "ALARM включен"],
            [
                "T6",
                "Авто-закрытие дверей при DOOR1/2_AUTO_CLOSE_SEC>0",
                "Двери закрываются по таймеру",
            ],
            [
                "T7",
                "SENSOR_MODE=serial, строки D1:OPEN/CLOSED",
                "FSM получает DoorClosedChanged, статус обновляется",
            ],
            [
                "T8",
                "Перезапуск с EYEGATE_DB_PATH=data/eyegate_scud.db",
                "Пользователи сохраняются",
            ],
            [
                "T9",
                "Страницы SPA /kiosk /monitor /sim /admin /enroll открываются",
                "Маршруты работают, 404 нет",
            ],
            ["T10", "ENROLL: /admin Capture face", "has_face=true в users"],
        ],
    )

    doc.add_heading("3 Методики", level=1)
    doc.add_paragraph(
        "T1: Открыть /monitor, проверить сетевые запросы → только /api/video/mjpeg; выключить камеру — сообщение CAMERA DOWN."
    )
    doc.add_paragraph(
        "T2: Подключиться WS /ws/status, убедиться в первом сообщении со всеми полями GateStatus."
    )
    doc.add_paragraph(
        "T3: Через /kiosk ввести PIN demo → Door1 unlock, после закрытия Door1 и анализа → Door2 unlock."
    )
    doc.add_paragraph(
        "T7: SENSOR_MODE=serial, отправить строки через SerialBridge или эмуляцию; проверить snapshot."
    )

    doc.add_heading("4 Критерии приемки", level=1)
    doc.add_paragraph(
        "Все тесты T1..T10 выполнены успешно; проблемы A–F закрыты или имеют оформленный план внедрения."
    )

    doc.add_heading("5 Контрольные примеры (данные)", level=1)
    add_table(
        doc,
        ["Пример", "Вход", "Ожидаемый выход"],
        [
            ["C1", "PIN demo (CARD123)", "ACCESS_GRANTED, Door2 unlock"],
            ["C2", "Unknown PIN/лицо", "ACCESS_DENIED"],
            ["C3", "2 человека, двери закрыты", "ALARM"],
            ["C4", "Serial D1:OPEN", "door1_closed=false в GateStatus"],
        ],
    )

    doc.save(dest)


def main() -> None:
    PMI_DIR.mkdir(parents=True, exist_ok=True)
    facts = load_facts()
    build_pmi(facts)
    print("Built ПМИ/ПСИ документ (плейсхолдер).")


if __name__ == "__main__":
    main()
