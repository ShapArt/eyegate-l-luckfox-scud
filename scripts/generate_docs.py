from __future__ import annotations

import datetime as dt
import json
import os
import textwrap
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"

OUTPUT_RPZ = DOCS_DIR / "EyeGate_Mantrap_RPZ.docx"
OUTPUT_USER = DOCS_DIR / "EyeGate_Mantrap_USER_GUIDE.docx"
OUTPUT_OPERATOR = DOCS_DIR / "EyeGate_Mantrap_OPERATOR_GUIDE.docx"
OUTPUT_TEST_PLAN = DOCS_DIR / "EyeGate_Mantrap_TEST_PLAN.docx"
OUTPUT_TRACE_XLSX = DOCS_DIR / "TRACEABILITY_MATRIX.xlsx"


def find_template() -> Optional[Path]:
    candidates = [
        Path("/mnt/data/РПЗ_СБ_Щит (1).docx"),
        ROOT / "docs" / "templates" / "RPZ_template.docx",
    ]
    for cand in candidates:
        if cand.exists():
            return cand
    return None


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        body.remove(child)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def parse_env_example() -> List[Dict[str, str]]:
    path = ROOT / ".env.example"
    envs: List[Dict[str, str]] = []
    if not path.exists():
        return envs
    descriptions = {
        "EYEGATE_ENV": "Среда (dev/prod) для FastAPI",
        "EYEGATE_HOST": "Хост для сервера",
        "EYEGATE_PORT": "Порт FastAPI",
        "EYEGATE_DUMMY_HW": "Использовать симуляцию дверей/сирены",
        "VISION_MODE": "Режим vision: real | dummy | browser",
        "VISION_MATCH_THRESHOLD": "Порог совпадения лица",
        "VISION_TTL_SEC": "Время устаревания кадра",
        "VISION_MATCH_METRIC": "Метрика сравнения (l2/cosine)",
        "VISION_AUTO_DOWNLOAD": "Автозагрузка моделей YuNet/SFace",
        "VISION_DUMMY_PEOPLE": "Dummy: количество людей",
        "VISION_DUMMY_RECOGNIZED": "Dummy: ID распознанных",
        "VISION_DUMMY_DELAY_MS": "Dummy: задержка ответа, мс",
        "EYEGATE_ENTER_TIMEOUT": "Таймаут входа (до закрытия двери 1)",
        "EYEGATE_CHECK_TIMEOUT": "Таймаут анализа комнаты",
        "EYEGATE_EXIT_TIMEOUT": "Таймаут выхода (дверь 2)",
        "EYEGATE_ALARM_TIMEOUT": "Таймаут сигнализации",
        "ALLOW_MULTI_KNOWN": "Разрешить несколько известных",
        "REQUIRE_FACE_MATCH_FOR_DOOR2": "Требовать лицо для двери 2",
        "MAX_PEOPLE_ALLOWED": "Макс. людей в шлюзе",
        "DOOR1_CLOSE_STABILIZE_MS": "Задержка стабилизации после закрытия двери 1",
        "ROOM_CHECK_SAMPLES": "Количество выборок Vision перед решением",
        "DOOR_AUTO_CLOSE_SEC": "Автодоводчик обеих дверей (сек)",
        "DOOR1_AUTO_CLOSE_SEC": "Автодоводчик двери 1 (сек)",
        "DOOR2_AUTO_CLOSE_SEC": "Автодоводчик двери 2 (сек)",
        "SIM_AUTO_CLOSE_MS": "Наследный авто-close для симулятора (мс)",
        "SENSOR_MODE": "Источник датчиков: sim | serial",
        "SENSOR_SERIAL_PORT": "COM-порт для SerialBridge",
        "SENSOR_SERIAL_BAUD": "Скорость SerialBridge",
        "ADMIN_TOKEN": "Сервисный токен админа (header)",
        "AUTH_SECRET": "Секрет для JWT",
        "ADMIN_LOGIN": "Логин встроенного админа",
        "ADMIN_PASS": "Пароль встроенного админа",
        "ADMIN_CARD_ID": "Card ID встроенного админа",
        "EYEGATE_DEMO_MODE": "Демо-режим (расширенные права, временный DB)",
        "EYEGATE_DB_PATH": "Явный путь к SQLite",
    }
    with path.open(encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if "=" not in line:
                continue
            key, value = line.split("=", 1)
            envs.append(
                {
                    "key": key.strip(),
                    "value": value.strip(),
                    "description": descriptions.get(key.strip(), ""),
                }
            )
    return envs


def collect_routes() -> List[Dict[str, Any]]:
    from fastapi.routing import APIRoute, WebSocketRoute

    from server.main import create_app

    app = create_app()
    routes: List[Dict[str, Any]] = []
    for route in app.routes:
        if isinstance(route, APIRoute):
            methods = sorted(m for m in route.methods if m not in {"HEAD", "OPTIONS"})
            if not methods:
                continue
            if route.path.startswith(("/openapi", "/docs", "/redoc")):
                continue
            description = (
                (route.summary or "").strip()
                or (route.description or "").strip()
                or ((route.endpoint.__doc__ or "").strip() if route.endpoint else "")
            )
            routes.append(
                {
                    "path": route.path,
                    "methods": methods,
                    "name": getattr(route.endpoint, "__name__", ""),
                    "description": description,
                }
            )
        elif isinstance(route, WebSocketRoute):
            routes.append(
                {
                    "path": route.path,
                    "methods": ["WEBSOCKET"],
                    "name": getattr(route.endpoint, "__name__", ""),
                    "description": "WebSocket endpoint",
                }
            )
    return routes


def collect_db_schema() -> Dict[str, Any]:
    from db.base import get_connection, get_db_path
    from db.init_db import init_db

    init_db()
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("SELECT name FROM sqlite_master WHERE type='table';")
    tables: Dict[str, Any] = {}
    for row in cur.fetchall():
        name = row["name"] if hasattr(row, "keys") else row[0]
        cur.execute(f"PRAGMA table_info({name});")
        cols = [
            {
                "name": col["name"],
                "type": col["type"],
                "notnull": bool(col["notnull"]),
                "pk": bool(col["pk"]),
                "default": col["dflt_value"],
            }
            for col in cur.fetchall()
        ]
        tables[name] = cols
    return {"path": str(get_db_path()), "tables": tables}


def collect_snapshot() -> Dict[str, Any]:
    from server.deps import get_gate_controller

    controller = get_gate_controller()
    snap = controller.snapshot()

    def _normalize(obj: Any) -> Any:
        if isinstance(obj, dt.datetime):
            return obj.isoformat()
        if isinstance(obj, (list, tuple)):
            return [_normalize(v) for v in obj]
        if isinstance(obj, dict):
            return {k: _normalize(v) for k, v in obj.items()}
        return obj

    return _normalize(snap)


def collect_fsm() -> Dict[str, Any]:
    from gate.models import GateState

    transitions = [
        {
            "state": "IDLE",
            "event": "CardPresented/AuthResult(allow)",
            "next": "WAIT_ENTER",
            "actions": "Unlock door1, start enter timeout",
        },
        {
            "state": "WAIT_ENTER",
            "event": "Door1 closed",
            "next": "CHECK_ROOM",
            "actions": "Lock door1, lock door2, start room analysis + check timeout",
        },
        {
            "state": "WAIT_ENTER",
            "event": "Timeout enter",
            "next": "IDLE",
            "actions": "Lock door1, cancel timeouts, clear ctx",
        },
        {
            "state": "CHECK_ROOM",
            "event": "Policy open_door2",
            "next": "ACCESS_GRANTED",
            "actions": "Lock door1, unlock door2, start exit timeout",
        },
        {
            "state": "CHECK_ROOM",
            "event": "Policy alarm",
            "next": "ALARM",
            "actions": "Lock both, alarm on, alarm timeout",
        },
        {
            "state": "CHECK_ROOM",
            "event": "Policy deny",
            "next": "ACCESS_DENIED",
            "actions": "Lock both, cancel timeouts",
        },
        {
            "state": "CHECK_ROOM",
            "event": "Timeout check",
            "next": "IDLE",
            "actions": "Lock both, cancel timeouts, clear ctx",
        },
        {
            "state": "ACCESS_GRANTED",
            "event": "Door2 closed",
            "next": "IDLE",
            "actions": "Lock door2, cancel exit timeout, clear ctx",
        },
        {
            "state": "ACCESS_GRANTED",
            "event": "Timeout exit",
            "next": "RESET",
            "actions": "Lock door2, cancel exit timeout, clear ctx",
        },
        {
            "state": "ALARM",
            "event": "Reset/timeout",
            "next": "RESET",
            "actions": "Alarm off, lock both, cancel timeouts",
        },
        {
            "state": "RESET",
            "event": "Any",
            "next": "IDLE",
            "actions": "Log and clear ctx",
        },
    ]
    return {"states": [s.name for s in GateState], "transitions": transitions}


def collect_repo_facts() -> Dict[str, Any]:
    facts: Dict[str, Any] = {}
    facts["template"] = find_template()
    facts["env_vars"] = parse_env_example()
    facts["routes"] = collect_routes()
    facts["db"] = collect_db_schema()
    facts["snapshot"] = collect_snapshot()
    facts["fsm"] = collect_fsm()
    facts["tree"] = {
        "server": [p.name for p in (ROOT / "server").iterdir()],
        "gate": [p.name for p in (ROOT / "gate").iterdir()],
        "vision": [p.name for p in (ROOT / "vision").iterdir()],
        "hw": [p.name for p in (ROOT / "hw").iterdir()],
        "db": [p.name for p in (ROOT / "db").iterdir()],
        "web/app/src": [p.name for p in (ROOT / "web" / "app" / "src").iterdir()],
        "tests": [p.name for p in (ROOT / "tests").iterdir()],
        "docs": [p.name for p in (ROOT / "docs").iterdir()],
    }
    facts["entrypoint"] = (
        "server/main.py:create_app() mounts /api, /ws/status, SPA catch-all"
    )
    facts["video"] = {
        "mjpeg": "/api/video/mjpeg",
        "snapshot": "/api/video/snapshot",
        "ws": "/ws/status",
        "refresh_sec": 0.08,
        "legacy_monitor_html": "web/templates/monitor.html uses getUserMedia (browser camera)",
        "react_monitor": "web/app/src/pages/MonitorPage.tsx uses <img src=/api/video/mjpeg>",
    }
    facts["vision"] = {
        "people_counter": "vision/people_count.py background subtraction silhouettes",
        "face_detection": "vision/service.py YuNet or Haar; SFace recognizer if models exist",
        "match_threshold": os.getenv("VISION_MATCH_THRESHOLD", "0.6"),
        "dummy": "VisionServiceDummyControl with env VISION_DUMMY_*",
        "outputs": [
            "people_count",
            "faces[].label/user_id/score",
            "match/matched_user_id",
            "recognized_user_ids/scores",
        ],
    }
    facts["known_issues"] = [
        "Legacy /monitor template uses браузерную камеру (face-api.js) вместо бэкенд /api/video/mjpeg; React-страница уже использует MJPEG.",
        "Файлы SQLite в data/ пустые и создаются fallback_* — при блокировке БД init_db делает новый путь, что выглядит как потеря данных после рестартов.",
        "Автодоводчик реализован в hw/simulated.py и конфиг env DOOR*_AUTO_CLOSE_SEC, но в GPIO-контроллере нет таймера — требуется вынести в реальный режим.",
        "UI должен показывать имя/UNKNOWN (красный), не Face 1 — MonitorPage делает это через API users; legacy overlay в vision/service.py подписывает UNKNOWN, требуется выравнивание.",
        'Для Proteus/SerialBridge нужен устоявшийся протокол D1:OPEN/D2:CLOSED или JSON {"door":1,"state":"open"} и инструкция подключения.',
    ]
    return facts


def paragraph(doc: Document, text: str, style: Optional[str] = None) -> None:
    for part in text.strip().split("\n"):
        doc.add_paragraph(part.strip(), style=style)


def add_title(
    doc: Document,
    title: str,
    subtitle: str,
    city: str = "Город",
    year: str = str(dt.datetime.now().year),
) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(title).bold = True
    doc.add_paragraph(subtitle).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Выполнил: __________________________").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("Группа: ____________________________").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph("Руководитель: _______________________").alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    doc.add_paragraph(f"{city}, {year}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def table_from_rows(doc: Document, headers: List[str], rows: List[List[Any]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        hdr[idx].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            cells[idx].text = "" if val is None else str(val)


def build_rpz(doc: Document, facts: Dict[str, Any]) -> None:
    add_title(doc, "EyeGate Mantrap", "Расчётно-пояснительная записка")
    doc.add_heading("СОДЕРЖАНИЕ", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("Введение", level=1)
    paragraph(
        doc,
        textwrap.dedent(
            """
            EyeGate Mantrap — учебно-демонстрационный комплекс шлюзовой зоны с двумя дверями, камерой и анти-tailgating логикой.
            Проект состоит из FastAPI бэкенда и React/Vite SPA фронтенда, поддерживает веб-сокеты статуса, MJPEG видеопоток и SQLite.
            Цель документа — зафиксировать требования, архитектуру, алгоритмы, текущие ограничения и программу испытаний.
            """
        ),
    )

    doc.add_heading("1 Назначение и область применения", level=1)
    doc.add_heading("1.1 Цели разработки", level=2)
    paragraph(
        doc,
        "Обеспечить контролируемый доступ через шлюз с раздельным управлением дверями, проверкой наличия одиночного пользователя и распознаванием лица.",
    )
    doc.add_heading("1.2 Основные функции", level=2)
    paragraph(
        doc,
        "- авторизация по PIN/картам через API /api/auth/*;\n"
        "- управление дверьми через FSM (gate/fsm.py) и симулятор hw/simulated.py;\n"
        "- видео: /api/video/mjpeg, /api/video/snapshot, overlay статуса в /monitor;\n"
        "- учёт событий в SQLite (db/models.py) и веб-сокет /ws/status.",
    )

    doc.add_heading("2 Исходные данные и допущения", level=1)
    paragraph(
        doc,
        "Исходные данные: код репозитория, .env.example, схемы FSM (gate/fsm.py), политика (policy/access.py), UI страницы /kiosk /monitor /sim /admin. Допущения: нет внешних API, камера одна, датчики дверей доступны через симуляцию или SerialBridge.",
    )

    doc.add_heading("3 Нормативные ссылки и термины", level=1)
    paragraph(
        doc,
        "- ГОСТ 34.602-89 Техническое задание на АС;\n- ГОСТ 19.201-78 Пояснительная записка;\n- OWASP ASVS (базовые меры); терминология FastAPI/WS/MJPEG/FSM.",
    )

    doc.add_heading("4 Обзор и постановка задачи", level=1)
    paragraph(
        doc,
        "Шлюз (mantrap) имеет две двери. Пользователь предъявляет карту/PIN, проходит через дверь 1, камера/vision проверяет, что внутри один человек и лицо соответствует. Дверь 2 открывается или поднимается тревога. Требуется исключить tailgating и обеспечить аудит событий.",
    )

    doc.add_heading("5 Аппаратная схема и состав", level=1)
    paragraph(
        doc,
        "Текущая поставка использует симулятор дверей hw/simulated.py и dummy alarm. Плановая целевая платформа — Luckfox Pico Ultra: GPIO (hw/doors.py) для замков и концевиков, камера CSI/USB, SerialBridge для датчиков. Две двери, два замка, два геркона/датчика, сирена.",
    )

    doc.add_heading("6 Алгоритм работы и конечный автомат", level=1)
    paragraph(
        doc,
        "FSM описан в gate/fsm.py и обслуживается gate/controller.py. Состояния: "
        + ", ".join(facts["fsm"]["states"]),
    )
    table_from_rows(
        doc,
        ["Состояние", "Событие", "Переход", "Действия"],
        [
            [t["state"], t["event"], t["next"], t["actions"]]
            for t in facts["fsm"]["transitions"]
        ],
    )
    paragraph(
        doc,
        "Диаграмма (Mermaid, обновляется вручную при изменении кода):\n"
        "```mermaid\n"
        "stateDiagram-v2\n"
        "    IDLE --> WAIT_ENTER : Card/Auth OK\n"
        "    WAIT_ENTER --> CHECK_ROOM : Door1 closed\n"
        "    WAIT_ENTER --> IDLE : timeout_enter\n"
        "    CHECK_ROOM --> ACCESS_GRANTED : policy open_door2\n"
        "    CHECK_ROOM --> ACCESS_DENIED : policy deny\n"
        "    CHECK_ROOM --> ALARM : policy alarm\n"
        "    CHECK_ROOM --> IDLE : timeout_check\n"
        "    ACCESS_GRANTED --> IDLE : Door2 closed\n"
        "    ACCESS_GRANTED --> RESET : timeout_exit\n"
        "    ALARM --> RESET : reset/timeout\n"
        "    RESET --> IDLE : any\n"
        "```",
    )

    doc.add_heading("7 Программная архитектура", level=1)
    paragraph(
        doc,
        textwrap.dedent(
            """
            Бэкенд: FastAPI, точка входа server/main.py (create_app), маршруты router.py. SPA статически монтируется из web/app/dist; legacy HTML в web/templates/*. WS: /ws/status рассылает snapshots из GateController.
            Фронтенд: React/Vite, страницы /kiosk (PIN), /monitor (MJPEG + overlay), /sim (датчики/замки), /admin (пользователи, снимок), /enroll.
            Последовательность: PIN → GateController.AuthResult → unlock door1 → закрытие → room analysis (vision) → policy (policy/access.py) → unlock door2 или alarm → event log → WS broadcast → UI.
            """
        ),
    )

    doc.add_heading("8 Видеоподсистема и мониторинг", level=1)
    paragraph(
        doc,
        textwrap.dedent(
            f"""
            Потоки: {facts['video']['mjpeg']} (StreamingResponse, ~12 fps), /api/video/snapshot. VisionServiceOpenCV захватывает кадры, строит JPEG и overlay. React MonitorPage использует <img src="/api/video/mjpeg"> через компонент MjpegStream.
            Проблема: legacy /monitor (web/templates/monitor.html) использует getUserMedia/face-api.js и показывает браузерную камеру. Правильно — отображать бекенд MJPEG и серверные боксы/статус. План: удалить legacy шаблон или заменить <img src="/api/video/mjpeg">.
            Overlay: boxes и labels приходят из WS (vision.faces[].label или name из /api/users/), UNKNOWN подсвечивается красным в React.
            """
        ),
    )

    doc.add_heading("9 Vision: силуэт + лицо", level=1)
    paragraph(
        doc,
        textwrap.dedent(
            """
            Силуэт: vision/people_count.py, MOG2 + морфология, фильтр площади/аспекта. Лицо: vision/service.py — YuNet или Haar каскад, SFace для эмбеддингов; best_match_for_embedding с порогом VISION_MATCH_THRESHOLD (по умолчанию 0.6, метрика l2/cosine).
            Результат: people_count, faces[label/user_id/score/is_known], match/matched_user_id, recognized_user_ids/scores, frame_w/h. UNKNOWN подсвечивается на фронтенде красным; при совпадении выводится имя/логин.
            Отсутствует: полноценное хранение эмбеддингов кроме dummy/серого плейсхолдера; при stale кадрах FSM переводит в CHECK_ROOM повторно.
            """
        ),
    )

    doc.add_heading("10 База данных", level=1)
    paragraph(doc, f"Путь SQLite: {facts['db']['path']}")
    for tbl, cols in facts["db"]["tables"].items():
        table_from_rows(
            doc,
            [f"Таблица {tbl}", "Тип", "NOT NULL", "PK", "Default"],
            [[c["name"], c["type"], c["notnull"], c["pk"], c["default"]] for c in cols],
        )
    paragraph(
        doc,
        textwrap.dedent(
            """
            Наблюдение: в каталоге data/ множество файлов eyegate_scud_fallback_*.db нулевого размера — init_db переносит БД на новый путь при SQLITE_LOCKED. Это выглядит как потеря пользователей после рестартов.
            План фиксации: задать EYEGATE_DB_PATH на стабильный путь, проверить блокировки, убрать автоматический fallback на каждый запуск, добавить health-check на PRAGMA journal_mode=WAL и тест восстановления подключения.
            """
        ),
    )

    doc.add_heading("11 Интеграция с датчиками", level=1)
    paragraph(
        doc,
        textwrap.dedent(
            """
            Симуляция: hw/simulated.py, API /api/sim/* позволяет открыть/закрыть двери, управлять сенсорами и автодоводчиком (set_auto_close).
            Реал: hw/doors.py (GPIO, активный высокий/низкий), hw/alarm.py. SerialBridge (hw/serial_bridge.py) парсит строки D1:OPEN/D2:CLOSED или JSON {"door":1,"state":"open"} и вызывает on_sensor.
            Требуется Proteus-ready режим: задокументировать скорость, порт, формат; при SENSOR_MODE=serial и SENSOR_SERIAL_PORT=COM5 мост активируется из server.deps.start_gate_controller.
            Автодоводчик: в симуляторе работает по таймеру, в GPIO-режиме таймер отсутствует — нужно перенести логику в GateController или отдельный планировщик.
            """
        ),
    )

    doc.add_heading("12 Надёжность и отказоустойчивость", level=1)
    table_from_rows(
        doc,
        ["Отказ", "Последствие", "Мера/FMEA"],
        [
            [
                "Камера недоступна",
                "Vision stale → CHECK_ROOM зациклен",
                "Фолбэк dummy snapshot, запрет двери 2, тревога ALARM по политике",
            ],
            [
                "WS разрыв",
                "UI не обновляет статус",
                "Повторное подключение, REST /api/status/ как резерв",
            ],
            [
                "DB locked",
                "Потеря персистентности",
                "Явный EYEGATE_DB_PATH, busy_timeout=5000, убрать frequent fallback",
            ],
            [
                "Датчик завис в OPEN",
                "FSM ALARM (policy alarm)",
                "Дебаунс + ручной RESET оператором",
            ],
            [
                ">1 человек",
                "ALARM",
                "policy/access.py ALARM при people_count>max_people_allowed",
            ],
            [
                "Ложное лицо",
                "Неверный доступ",
                "Повысить VISION_MATCH_THRESHOLD, требовать match_distance<=0.6",
            ],
        ],
    )

    doc.add_heading("13 Безопасность", level=1)
    paragraph(
        doc,
        "- Угрозы: подмена камеры (неверный MJPEG), подмена датчиков (SerialBridge), подбор PIN, подмена лица/фото.\n"
        "- Меры: rate-limit логина (auth/rate_limit.py), ADMIN_TOKEN/Authorization Bearer, журнал событий db/logger.py, политика require_face_match_for_door2, deny/alarm при неизвестном лице, запрет открытия двери 2 без vision.",
    )

    doc.add_heading("14 Программа и методика испытаний (кратко)", level=1)
    paragraph(
        doc,
        "Полный ПМИ вынесен в отдельный файл EyeGate_Mantrap_TEST_PLAN.docx. Ключевые тесты: известный пользователь → Door2 open; неизвестный → ALARM; двое → ALARM; камера down → deny; проверка персистентности SQLite после restart; SerialBridge парсинг строк D1:OPEN.",
    )

    doc.add_heading("15 Технико-экономическое обоснование", level=1)
    paragraph(
        doc,
        "Используются бесплатные компоненты (FastAPI, React, SQLite, OpenCV). Железо: Luckfox Pico Ultra + камера + 2 электрозамка + 2 датчика — бюджетный комплект. Экономия за счёт симулятора для отладки без железа.",
    )

    doc.add_heading("16 Охрана труда / техника безопасности", level=1)
    paragraph(
        doc,
        "При работе с замками соблюдать требования электробезопасности, отключать питание перед подключением. Камера не должна нарушать приватность — хранить минимально необходимые данные, использовать стендовые данные для тестов.",
    )

    doc.add_heading("Заключение", level=1)
    paragraph(
        doc,
        "Проект реализует базовый шлюз с разделением дверей и минимальной политикой. Требуются доработки: стабильная БД, перевод мониторинга на серверный MJPEG, автодоводчик в GPIO-режиме, чёткий протокол SerialBridge.",
    )

    doc.add_heading("Список источников", level=1)
    paragraph(
        doc,
        "Код репозитория; FastAPI docs; OpenCV docs; ГОСТ 34.602-89; ГОСТ 19.201-78.",
    )

    doc.add_heading("Приложение A. Спецификация API", level=1)
    table_from_rows(
        doc,
        ["Метод", "Путь", "Назначение"],
        [
            [",".join(r["methods"]), r["path"], r["description"]]
            for r in facts["routes"]
        ],
    )

    doc.add_heading("Приложение B. Формат WS /ws/status", level=1)
    paragraph(doc, json.dumps(facts["snapshot"], ensure_ascii=False, indent=2))

    doc.add_heading("Приложение C. Переменные окружения", level=1)
    table_from_rows(
        doc,
        ["KEY", "Default", "Описание"],
        [[e["key"], e["value"], e["description"]] for e in facts["env_vars"]],
    )

    doc.add_heading("Приложение D. Матрица трассируемости (сводка)", level=1)
    paragraph(doc, "Полная версия в TRACEABILITY_MATRIX.xlsx.")

    doc.add_heading("Приложение E. Диаграммы", level=1)
    paragraph(
        doc,
        "FSM см. выше; sequence: PIN → door1 unlock → door1 close → vision analyze → policy → door2 unlock/alarm → WS/status → UI.",
    )


def build_user_guide(doc: Document, facts: Dict[str, Any]) -> None:
    add_title(doc, "EyeGate Mantrap", "Руководство пользователя")
    doc.add_heading("Назначение", level=1)
    paragraph(doc, "Простой сценарий доступа через /kiosk и наблюдение через /monitor.")

    doc.add_heading("Kiosk (/kiosk)", level=1)
    paragraph(
        doc,
        "1) Введите 4-значный PIN. 2) При успехе Door1 откроется (симулятор автоподнимает дверь). 3) Закройте Door1, дождитесь анализа камеры. 4) При успешном матчинге откроется Door2. При ALARM ждите оператора.",
    )

    doc.add_heading("Monitor (/monitor)", level=1)
    paragraph(
        doc,
        textwrap.dedent(
            f"""
            Поток: {facts['video']['mjpeg']}. При обрыве — кнопка Reload stream.
            Overlay: зелёный — известный, красный UNKNOWN. Значения People/FPS/Match отображаются в карточках.
            Legacy /monitor (HTML) использовать нельзя — он снимает браузерную камеру. Правильный монитор — React-страница.
            """
        ),
    )

    doc.add_heading("Simulator (/sim)", level=1)
    paragraph(
        doc,
        "Используйте Open/Close для дверей 1/2, наблюдайте сенсоры и замки. Для автодоводчика используйте API /api/sim/auto_close с delay_ms.",
    )

    doc.add_heading("Частые проблемы", level=1)
    paragraph(
        doc,
        "- ALARM сразу после входа: камера не видит лицо или два силуэта; убедитесь, что один человек.\n"
        "- Нет сохранения пользователей: проверьте EYEGATE_DB_PATH и права записи; очистите старые fallback БД.\n"
        "- Видео нет: убедитесь, что /api/video/mjpeg отдаёт JPEG; при dummy режим показывает плейсхолдер.",
    )


def build_operator_guide(doc: Document, facts: Dict[str, Any]) -> None:
    add_title(doc, "EyeGate Mantrap", "Руководство оператора/админа")
    doc.add_heading("Вход в систему", level=1)
    paragraph(
        doc,
        "Админский логин через /admin (UI) или POST /api/auth/admin/login с ADMIN_LOGIN/ADMIN_PASS. В демо режиме (EYEGATE_DEMO_MODE=1) доступ открыт.",
    )

    doc.add_heading("Создание пользователей (/admin)", level=1)
    paragraph(
        doc,
        "1) Заполните login, PIN, Name. 2) Нажмите Create — запись попадёт в users. 3) Нажмите Capture face для enrollment через backend snapshot (/api/video/snapshot).",
    )

    doc.add_heading("Мониторинг и реагирование", level=1)
    paragraph(
        doc,
        "- Следите за WS /ws/status и карточками MonitorPage.\n"
        "- При ALARM нажмите Reset (/api/status/reset) или кнопку в UI, убедитесь, что дверцы закрыты.\n"
        "- Экспорт событий: /api/events/export?format=csv|json (требует X-Admin-Token или Bearer admin).",
    )

    doc.add_heading("Работа с симулятором и датчиками", level=1)
    paragraph(
        doc,
        'Серийный мост: SENSOR_MODE=serial, SENSOR_SERIAL_PORT=COM5, формат D1:OPEN/D2:CLOSED или JSON {"door":1,"state":"open"}. Протестируйте через tests/test_serial_bridge.py.',
    )

    doc.add_heading("Обслуживание БД", level=1)
    paragraph(
        doc,
        "Установите EYEGATE_DB_PATH на постоянный путь (например, data/eyegate_scud.db). При замене БД остановите процесс, удалите пустые fallback_*.db, перезапустите init_db.",
    )


def build_test_plan(doc: Document, facts: Dict[str, Any]) -> None:
    add_title(doc, "EyeGate Mantrap", "Программа и методика испытаний")
    doc.add_heading("Общие сведения", level=1)
    paragraph(
        doc,
        "Тесты выполняются на FastAPI приложении, Vision в dummy или real режиме, БД SQLite.",
    )

    doc.add_heading("Тест-кейсы", level=1)
    cases = [
        [
            "TC1",
            "Известный пользователь 1 человек → Door2 open",
            "PIN login demo, закрыть Door1, ожидать ACCESS_GRANTED",
        ],
        [
            "TC2",
            "Неизвестный/нет лица → ALARM/ACCESS_DENIED",
            "people_count=1, match=False",
        ],
        ["TC3", "Два человека → ALARM", "people_count=2, policy/alarm"],
        [
            "TC4",
            "Камера недоступна → запрет открытия",
            "vision_error, stale snapshot => CHECK_ROOM повтор, нет unlock door2",
        ],
        [
            "TC5",
            "Персистентность БД после рестарта",
            "Создать пользователя, очистить cache controller, перечитать /api/users/",
        ],
        ["TC6", "SerialBridge парсинг", "D1:OPEN, D2:CLOSED строки → sensor callbacks"],
        [
            "TC7",
            "Автодоводчик",
            "set_auto_close delay_ms=2000 → дверь закрывается через ~2с",
        ],
        [
            "TC8",
            "Экспорт событий",
            "POST /api/status/reset; GET /api/events/export?format=csv",
        ],
        [
            "TC9",
            "WS broadcast",
            "Подписка на /ws/status, сравнить с /api/status/ snapshot",
        ],
    ]
    table_from_rows(doc, ["ID", "Цель", "Шаги/Ожидание"], cases)

    doc.add_heading("Критерии приёмки", level=1)
    paragraph(
        doc,
        "- Нет ложного открытия Door2 при people_count>1 или unknown лице;\n"
        "- Пользователи сохраняются между рестартами;\n"
        "- WS события соответствуют REST статусу;\n"
        "- MJPEG поток доступен с задержкой <1с.",
    )


def build_traceability(facts: Dict[str, Any]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Traceability"
    ws.append(["Требование", "Реализация (файлы)", "Тест/проверка"])
    rows = [
        [
            "PIN → Door1 unlock",
            "server/api/auth.py; gate/controller.py; web/app/src/pages/KioskPage.tsx",
            "tests/test_demo_api.py",
        ],
        [
            "Vision match для Door2",
            "vision/service.py; policy/access.py; gate/fsm.py",
            "tests/test_controller_policy_integration.py",
        ],
        [
            "WS статус /ws/status",
            "server/ws.py; server/status_bus.py",
            "tests/test_demo_api.py (snapshot)",
        ],
        [
            "MJPEG мониторинг",
            "server/api/video.py; web/app/src/pages/MonitorPage.tsx",
            "manual: open /monitor React",
        ],
        [
            "Автодоводчик",
            "hw/simulated.py set_auto_close; server/api/sim.py",
            "tests/test_api_sim.py",
        ],
        [
            "SerialBridge/Proteus",
            "hw/serial_bridge.py; server/deps.start_gate_controller",
            "tests/test_serial_bridge.py",
        ],
        [
            "Персистентность БД",
            "db/base.py; db/init_db.py",
            "tests/test_db_persistence.py",
        ],
        ["Экспорт событий", "server/api/events.py", "manual: GET /api/events/export"],
        [
            "ADMIN токен/логин",
            "server/deps.require_admin; server/api/auth.py",
            "manual: POST /api/auth/admin/login",
        ],
        [
            "UI имя/UNKNOWN",
            "vision/service.py labels; web/app/src/pages/MonitorPage.tsx",
            "manual: monitor overlay",
        ],
    ]
    for r in rows:
        ws.append(r)
    wb.save(OUTPUT_TRACE_XLSX)


def create_document(template: Optional[Path]) -> Document:
    if template:
        doc = Document(template)
        clear_body(doc)
        return doc
    return Document()


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    facts = collect_repo_facts()

    rpz = create_document(facts.get("template"))
    build_rpz(rpz, facts)
    rpz.save(OUTPUT_RPZ)

    user = create_document(facts.get("template"))
    build_user_guide(user, facts)
    user.save(OUTPUT_USER)

    operator = create_document(facts.get("template"))
    build_operator_guide(operator, facts)
    operator.save(OUTPUT_OPERATOR)

    test_plan = create_document(facts.get("template"))
    build_test_plan(test_plan, facts)
    test_plan.save(OUTPUT_TEST_PLAN)

    build_traceability(facts)

    summary = {
        "outputs": [
            str(OUTPUT_RPZ),
            str(OUTPUT_USER),
            str(OUTPUT_OPERATOR),
            str(OUTPUT_TEST_PLAN),
            str(OUTPUT_TRACE_XLSX),
        ],
        "template_used": str(facts.get("template")) if facts.get("template") else None,
        "routes": len(facts["routes"]),
        "db_path": facts["db"]["path"],
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
